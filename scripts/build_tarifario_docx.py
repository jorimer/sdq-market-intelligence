"""Genera el Tarifario PROVISIONAL v0.2 (DOCX) — artefacto de venta del PR-12 (Ola D).

Lee el catálogo de SKUs REAL de prod (data/artefactos_venta/skus_prod.json, dump de
GET /api/v1/billing/skus) y arma la propuesta de precios provisionales para revisión
del dueño. Los únicos números con decisión previa son el ancla del research custom
(US$3,500/encargo, PROVISIONAL — docs/MOTOR_RESEARCH_PILOTO.md) y el benchmark de la
consultoría tradicional (US$7,000–25,000+/encargo). El resto son PROPUESTA inicial:
nada rige hasta que el dueño lo publique en /admin/tarifario (tabla Tariff).

Uso: /opt/anaconda3/bin/python scripts/build_tarifario_docx.py
"""
import json
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x0F, 0x2A, 0x4A)
RED = RGBColor(0xD6, 0x2A, 0x2A)
GRAY = RGBColor(0x6B, 0x72, 0x80)

# Propuesta PROVISIONAL (USD) — para calibración del dueño; no publicada.
PROPOSED = {
    "insight_monthly": 149, "insight_annual": 1490,
    "deep_dive_once": 450,
    "all_access_monthly": 690, "all_access_annual": 6900,
    "enterprise_monthly": 1450, "enterprise_annual": 14500,
    "research_custom_once": 3500,  # ancla DECIDIDA (provisional) por el dueño
}


def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Inter"
    st.font.size = Pt(10)


def _h(doc, text, size=16, color=NAVY, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Plus Jakarta Sans"
    return p


def _p(doc, text, size=10, color=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.italic = italic
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(9)
    return t


def main():
    skus = json.load(open("data/artefactos_venta/skus_prod.json"))
    rows = skus.get("skus") or skus
    sectors = [r for r in rows if r["sku"].startswith("insight:")]

    doc = Document()
    _style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("SDQ·MIP — Tarifario PROVISIONAL")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY
    r.font.name = "Plus Jakarta Sans"
    _p(doc, "Versión 0.2 · " + date.today().strftime("%d/%m/%Y") +
       " · PROPUESTA PARA APROBACIÓN — ningún precio rige hasta publicarse en /admin/tarifario",
       size=9, color=RED, italic=True)

    _p(doc, "Catálogo comercial de SDQ Market Intelligence Platform: inteligencia "
            "financiera y sectorial de la República Dominicana con dato oficial "
            "(SIB, BCRD, SIPEN, SIS, ONE, DGA) y análisis generado por IA con "
            "guardrail anti-alucinación. Tres niveles por sector: Pulse (abierto), "
            "Insight (suscripción) y Deep Dive (informe a demanda).", size=10)

    _h(doc, "1 · Insight por sector — suscripción", 13)
    _p(doc, "Entidad/segmento nombrado: score y outlook, radar de pilares, pares, "
            "narrativa ejecutiva y alertas. Facturación mensual o anual (2 meses de "
            "descuento en el plan anual).", size=9, color=GRAY)
    sector_rows = [
        (s["label"].replace("Insight · ", ""),
         f"US$ {PROPOSED['insight_monthly']}/mes",
         f"US$ {PROPOSED['insight_annual']}/año")
        for s in sectors
    ]
    _table(doc, ["Sector (SKU insight:*)", "Mensual (propuesto)", "Anual (propuesto)"], sector_rows)

    _h(doc, "2 · Deep Dive — informe a demanda", 13)
    _p(doc, "Informe de calificación completa de UNA entidad (p. ej. un banco): todo "
            "el Insight + entorno operativo, escenarios/sensibilidades, evaluación de "
            "riesgo y recomendación. Entrega PDF/DOCX con marca SDQ.", size=9, color=GRAY)
    _table(doc, ["Producto", "Precio por informe (propuesto)"],
           [("Deep Dive · cualquier sector del catálogo", f"US$ {PROPOSED['deep_dive_once']}"),
            ("Muestras adjuntas", "Banreservas · Banco Múltiple Lafise (2026-Q1)")])

    _h(doc, "3 · Bundles", 13)
    _table(doc, ["Bundle", "Mensual (propuesto)", "Anual (propuesto)", "Incluye"],
           [("All-Access", f"US$ {PROPOSED['all_access_monthly']}", f"US$ {PROPOSED['all_access_annual']}",
             "Todos los Insight sectoriales"),
            ("Enterprise", f"US$ {PROPOSED['enterprise_monthly']}", f"US$ {PROPOSED['enterprise_annual']}",
             "Todo el catálogo (Insight + Deep Dives)")])

    _h(doc, "4 · Research Custom — síntesis cross-dominio", 13)
    _p(doc, "Pregunta libre del cliente → el motor convoca los ejes pertinentes "
            "(banca, macro, monetario, comercio, pensiones, seguros…) y emite un "
            "dictamen integrado con evidencia por eje. Referencia de mercado de la "
            "consultoría tradicional: US$7,000–25,000+ por encargo, con 2–4 semanas "
            "de espera.", size=9, color=GRAY)
    _table(doc, ["Producto", "Precio (provisional decidido)"],
           [("special:research-custom · dictamen a medida", f"US$ {PROPOSED['research_custom_once']} por encargo")])

    _h(doc, "5 · Notas", 13)
    for n in (
        "Los precios son PROVISIONALES (propuesta v0.2). Rigen únicamente al publicarse "
        "en la tabla de tarifas (/admin/tarifario), con vigencia y moneda explícitas.",
        "Impuestos: ITBIS 18% se agrega al total según residencia fiscal (exento "
        "exterior); el desglose Subtotal + ITBIS = Total se muestra en el checkout.",
        "Pago self-serve vía PayPal (tarjeta o cuenta); factura PDF automática con "
        "desglose fiscal. e-CF/DGII listo para habilitarse.",
        "Pulse permanece abierto (agregado del sistema, sin nombrar entidades) como "
        "canal de adquisición.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(n)
        r.font.size = Pt(9)

    out = "data/artefactos_venta/SDQ_MIP_Tarifario_PROVISIONAL_v0.2.docx"
    doc.save(out)
    print("OK", out)


if __name__ == "__main__":
    main()
