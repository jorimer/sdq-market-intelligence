"""Genera el Catálogo SDQMIP v3 (estado REAL) + one-pager de estado para reunión.

Reemplaza como fuente comercial al Catalogo_SDQMIP_v2 (21-jun-2026), que quedó
congelado pre-productización: describía 1 solo eje vendible y 9 en roadmap. El
estado real verificado contra PRODUCCIÓN (2026-07-13): 14 sectores publicados,
cada uno con Pulse · Insight · Deep Dive, tarifario vigente (provisional v0.2) y
checkout self-serve PayPal operativo.

Fuentes de datos: shared.products.registry (catálogo/fuentes por sector, el mismo
que sirve la plataforma) + data/artefactos_venta/skus_prod.json (precios vigentes).

Uso: /opt/anaconda3/bin/python scripts/build_catalogo_v3.py
"""
import sys
from datetime import date

sys.path.insert(0, ".")

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

from shared.products.registry import PRODUCT_CATALOG  # noqa: E402

NAVY = RGBColor(0x0F, 0x2A, 0x4A)
RED = RGBColor(0xD6, 0x2A, 0x2A)
GRAY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x11, 0x7A, 0x3D)

OUT_DIR = "/Users/ricardomercado/Developer/SDQMIP"
TODAY = date.today().strftime("%d de %B de %Y").replace("July", "julio")

PRICING_NOTE = (
    "Precios VIGENTES hoy en plataforma (tarifario provisional v0.2, publicado "
    "13-jul-2026): Insight US$149/mes · US$1,490/año por sector; Deep Dive US$450 "
    "por informe; All-Access US$690/mes · US$6,900/año; Enterprise US$1,450/mes · "
    "US$14,500/año; Research Custom US$3,500 por encargo (ancla decidida). "
    "DECISIÓN ABIERTA de dirección: estos montos siguen la filosofía self-serve; "
    "el estudio comercial (v2) anclaba en consultoría boutique "
    "(Insight US$4,000–15,000/año · Deep Dive US$3,500–20,000). Recalibrar el "
    "tarifario toma minutos (editor masivo + sync automático de planes PayPal)."
)


def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Inter"
    st.font.size = Pt(10)


def _h(doc, text, size=15, color=NAVY, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Plus Jakarta Sans"
    return p


def _p(doc, text, size=10, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.italic = italic
    r.bold = bold
    return p


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(9)
    return t


def sector_rows():
    return [
        (e.display_name.replace("SDQ ", ""), "LISTO — en producción",
         "Pulse · Insight · Deep Dive", e.source)
        for e in PRODUCT_CATALOG
    ]


def build_catalogo():
    doc = Document()
    _style(doc)

    r = doc.add_paragraph().add_run("SDQMIP — Catálogo de Productos Sectoriales")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = NAVY
    r.font.name = "Plus Jakarta Sans"
    _p(doc, f"v3 · {TODAY} · Documento comercial confidencial · Santo Domingo, RD",
       size=9, color=GRAY, italic=True)
    _p(doc, "REEMPLAZA al Catálogo v2 (21-jun-2026), que quedó desactualizado: "
            "describía un solo eje vendible y nueve en construcción. Estado de este "
            "documento verificado contra PRODUCCIÓN el 13-jul-2026 [Certain].",
       size=9, color=RED, italic=True)

    _h(doc, "Modelo de empaquetado — un producto por sector, tres niveles")
    _p(doc, "Pulse (abierto; agregado del sistema, sin nombrar entidades — canal de "
            "adquisición) · Insight (suscripción mensual/anual; entidad o segmento "
            "nombrado, score+outlook, pares, narrativa IA con guardrail numérico) · "
            "Deep Dive (informe a demanda: Insight + entorno operativo, escenarios, "
            "riesgo y recomendación; PDF/DOCX con marca SDQ). Diferenciador "
            "transversal intacto: cada número es auditable hasta su fuente "
            "autoritativa; lo que no tiene dato se declara (N/D), no se rellena.")
    _p(doc, PRICING_NOTE, size=9, color=GRAY)

    _h(doc, "Estado real del portafolio — 14 sectores vendibles hoy")
    _p(doc, "LISTO = publicado en producción con dato real ingerido, gate de "
            "readiness aprobado, precio vigente y descarga operativa. Los tres "
            "niveles de cada sector están cableados de punta a punta.", size=9, color=GRAY)
    _table(doc,
           ["Sector", "Estado", "Niveles", "Fuentes autoritativas (el foso)"],
           sector_rows())

    _h(doc, "Qué cambió desde el v2 (21-jun → 13-jul)")
    for txt in (
        "Los 9 sectores que el v2 listaba 'En desarrollo/Roadmap' (macro, comercio, "
        "turismo, zonas francas, energía, telecom, construcción, agroindustria, ESG) "
        "pasaron a LISTO: productización 10×3 completa con dato real y gates de "
        "publicación.",
        "Se sumaron 4 ejes que el v2 no contemplaba: Pensiones (SIPEN, ISA por AFP), "
        "Seguros (SIS: ISF por aseguradora · SISALRIL: cobertura SFS + ISARS por ARS "
        "— 'Salud' dejó de estar descartado), Política Monetaria (modelo TPM con "
        "track record en vivo) y Estructura de la Economía.",
        "Canal self-serve operativo: tarifario publicado (30 SKUs), checkout PayPal "
        "(compra puntual y suscripción) con ITBIS desglosado y factura automática; "
        "planes de PayPal se sincronizan solos desde el tarifario.",
        "Motor de Research Custom (síntesis cross-dominio, US$3,500/encargo "
        "provisional) construido y pilotado — decisión vigente: NO ofrecerlo aún.",
        "Muestras comerciales reales disponibles: Deep Dive de Banreservas "
        "(SDQ-AA · 86.7/100, 14 pp) y Banco Múltiple Lafise (SDQ-A+ · 78.0/100, "
        "15 pp), generados de producción con dato SIB 2026-Q1.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(txt)
        run.font.size = Pt(9.5)

    _h(doc, "Sectores evaluados y descartados")
    _p(doc, "Retail/Consumo (sin dato abierto de grado comercial en RD). Sector "
            "Público se pliega dentro de Macro & Country Risk. Salud SALIÓ de esta "
            "lista: se cubre vía SISALRIL dentro del eje Seguros.", size=9)

    out = f"{OUT_DIR}/Catalogo_SDQMIP_v3_Productos_Sectoriales.docx"
    doc.save(out)
    return out


def build_onepager():
    doc = Document()
    _style(doc)
    r = doc.add_paragraph().add_run("SDQMIP — Estado real de la plataforma")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    r.font.name = "Plus Jakarta Sans"
    _p(doc, f"Para reunión de equipo · {TODAY} · verificado contra producción hoy",
       size=9, color=GRAY, italic=True)

    _p(doc, "14 ejes sectoriales vendibles HOY, cada uno con Pulse · Insight · "
            "Deep Dive en producción — no 'uno listo y nueve en roadmap' (ese era "
            "el estado de junio).", bold=True, size=11)

    _table(doc, ["Sector", "Estado", "Fuentes"],
           [(n, e, f) for (n, e, _lv, f) in sector_rows()])

    _h(doc, "Comercial", 12)
    for txt in (
        "Tarifario vigente (provisional v0.2): Insight US$149/mes · 1,490/año por "
        "sector · Deep Dive US$450 · All-Access 690/mes · Enterprise 1,450/mes · "
        "Research Custom US$3,500/encargo.",
        "Checkout self-serve PayPal operativo (sandbox): compra puntual y "
        "suscripción, ITBIS desglosado, factura automática. Falta solo la "
        "aprobación buyer-sandbox para el E2E completo.",
        "Muestras reales: Deep Dive Banreservas (SDQ-AA · 86.7) y Lafise "
        "(SDQ-A+ · 78.0), dato SIB 2026-Q1.",
        "DECISIÓN ABIERTA: ancla de precios — self-serve (v0.2 vigente) vs "
        "consultoría boutique (estudio: Insight 4,000–15,000/año). Recalibrar "
        "toma minutos.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(txt)
        run.font.size = Pt(9.5)

    out = f"{OUT_DIR}/SDQMIP_Estado_Real_OnePager_2026-07-13.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_catalogo())
    print(build_onepager())
