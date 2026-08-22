"""Genera el Dossier Comercial y de Marca de SDQ·MIP (DOCX + PDF).

Artefacto de HANDOFF para un diseñador senior que construirá la presentación comercial
de la plataforma. Reúne en un solo documento (a) la identidad y los tokens de marca ya
vigentes en el repo, (b) el inventario completo de features del catálogo y de las
capacidades transversales, y (c) el modelo comercial con las reglas de lo que se puede
y no se puede afirmar.

Doctrina que este script respeta y que hay que seguir respetando al editarlo:

- **Ninguna cifra de validación se escribe a mano.** El documento NO transcribe Ginis,
  IC ni N: declara de dónde se leen (`GET /api/v1/products/credenciales`) y por qué una
  cifra copiada se desincroniza. Ver `docs/CLAIMS_COMERCIALES.md`.
- **El catálogo se lee del código**, no de una lista paralela: `PRODUCT_CATALOG` y
  `catalog_skus()` son la fuente. Un producto nuevo aparece acá sin editar este archivo.
- **Los precios son PROPUESTA** hasta que el dueño los publique en `/admin/tarifario`.

Assets: `docs/comercial/assets/` (generados aparte, ver el README de esa carpeta).

Uso:
    pip install python-docx
    python scripts/build_dossier_comercial.py [--out docs/comercial]
"""
from __future__ import annotations

import argparse
import os
import sys
from datetime import date
from typing import Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ── Tokens de marca (espejo de frontend/src/index.css — dirección "Claro & Vivo") ──
CANVAS = "F5F8FC"
SURFACE = "FFFFFF"
SURFACE2 = "F1F5FB"
INK = "0A1A3A"
BODY = "43506B"
MUTED = "76829C"
FAINT = "9AA6BF"
BORDER = "E7ECF3"
BORDER_STRONG = "D6DEEC"
ACCENT = "1E6FFF"
ACCENT_HOVER = "1A60E0"
ACCENT_SOFT = "EAF1FF"
ACCENT_INK = "1551C0"
TEAL = "0F7E7E"
TEAL_SOFT = "E3F2F2"
OK = "15875A"
OK_SOFT = "E5F3EC"
WARN = "B7791F"
WARN_SOFT = "FBF1DD"
ALERT = "C8392E"
ALERT_SOFT = "FBEAE8"
C1, C2, C3, C4, C5, C6 = "1E6FFF", "0F7E7E", "7A5AF8", "B7791F", "E0729A", "2BA8A8"
GRID = "EAEFF6"

DARK = {
    "--canvas": "0A0F1C", "--surface": "121B2E", "--surface-2": "182238",
    "--ink": "EAF0FB", "--body": "AEB9D2", "--muted": "6E7C9B", "--faint": "55617E",
    "--border": "232F49", "--border-strong": "2E3C5A",
    "--accent": "3B82F6", "--accent-hover": "5A96F8", "--accent-ink": "93BBFB",
    "--teal": "2DD4BF", "--ok": "34D399", "--warn": "FBBF24", "--alert": "F2645A",
    "--c1": "3B82F6", "--c2": "2DD4BF", "--c3": "A78BFA",
    "--c4": "FBBF24", "--c5": "F472B6", "--c6": "22D3EE", "--grid": "1E2940",
}

DISPLAY_FONT = "Plus Jakarta Sans"
BODY_FONT = "Inter"
MONO_FONT = "JetBrains Mono"


def rgb(hexstr: str) -> RGBColor:
    return RGBColor.from_string(hexstr)


# ─────────────────────────── primitivas de layout ───────────────────────────
#
# OOXML impone el ORDEN de los hijos de `pPr`, `tcPr`, `tblPr` y `rPr`: un elemento
# correcto en la posición equivocada invalida el archivo y LibreOffice se niega a
# abrirlo (se detectó así al convertir a PDF). Por eso nada se agrega con `append`:
# todo pasa por `_insert_ordered`, que respeta la secuencia del esquema.

_PPR_ORDER = (
    "pStyle keepNext keepLines pageBreakBefore framePr widowControl numPr "
    "suppressLineNumbers pBdr shd tabs suppressAutoHyphens kinsoku wordWrap "
    "overflowPunct topLinePunct autoSpaceDE autoSpaceDN bidi adjustRightInd "
    "snapToGrid spacing ind contextualSpacing mirrorIndents suppressOverlap jc "
    "textDirection textAlignment textboxTightWrap outlineLvl divId cnfStyle rPr "
    "sectPr pPrChange").split()

_TCPR_ORDER = (
    "cnfStyle tcW gridSpan hMerge vMerge tcBorders shd noWrap tcMar textDirection "
    "tcFitText vAlign hideMark headers cellIns cellDel cellMerge tcPrChange").split()

_TBLPR_ORDER = (
    "tblStyle tblpPr tblOverlap bidiVisual tblStyleRowBandSize tblStyleColBandSize "
    "tblW jc tblCellSpacing tblInd tblBorders shd tblLayout tblCellMar tblLook "
    "tblCaption tblDescription tblPrChange").split()

_TRPR_ORDER = (
    "cnfStyle divId gridBefore gridAfter wBefore wAfter cantSplit trHeight tblHeader "
    "tblCellSpacing jc hidden ins del trPrChange").split()

_RPR_ORDER = (
    "rStyle rFonts b bCs i iCs caps smallCaps strike dstrike outline shadow emboss "
    "imprint noProof snapToGrid vanish webHidden color spacing w kern position sz "
    "szCs highlight u effect bdr shd fitText vertAlign rtl cs em lang "
    "eastAsianLayout specVanish oMath rPrChange").split()

_ORDERS = {"pPr": _PPR_ORDER, "tcPr": _TCPR_ORDER, "tblPr": _TBLPR_ORDER,
           "trPr": _TRPR_ORDER, "rPr": _RPR_ORDER}


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _insert_ordered(parent, child):
    """Inserta ``child`` en ``parent`` en la posición que exige el esquema OOXML."""
    order = _ORDERS.get(_local(parent.tag))
    if order is None:
        parent.append(child)
        return child
    name = _local(child.tag)
    try:
        rank = order.index(name)
    except ValueError:
        parent.append(child)
        return child
    for existing in parent:
        ex = _local(existing.tag)
        if ex not in order or order.index(ex) > rank:
            existing.addprevious(child)
            return child
    parent.append(child)
    return child


def _ensure(parent, name: str):
    """Devuelve el hijo ``w:{name}`` de ``parent``, creándolo en su posición si falta."""
    found = parent.find(qn(f"w:{name}"))
    if found is None:
        found = _insert_ordered(parent, OxmlElement(f"w:{name}"))
    return found

def _set_font(run, *, font=BODY_FONT, size=10.0, color=BODY, bold=False,
              italic=False, caps=False, spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = _ensure(rpr, "rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    if caps:
        el = OxmlElement("w:caps")
        el.set(qn("w:val"), "1")
        _insert_ordered(rpr, el)
    if spacing is not None:  # en veinteavos de punto
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(int(spacing)))
        _insert_ordered(rpr, el)
    return run


def _shd(pr, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _insert_ordered(pr, shd)


def _shade(element, fill: str):
    """Sombreado de párrafo (recibe el elemento ``w:p``)."""
    _shd(element.get_or_add_pPr(), fill)


def _cell_shade(cell, fill: str):
    _shd(cell._tc.get_or_add_tcPr(), fill)


def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcpr = cell._tc.get_or_add_tcPr()
    old = tcpr.find(qn("w:tcMar"))
    if old is not None:
        tcpr.remove(old)
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    _insert_ordered(tcpr, mar)


def _tbl_borders(table, *, val="single", sz=4, color=BORDER):
    tblpr = table._tbl.tblPr
    old = tblpr.find(qn("w:tblBorders"))
    if old is not None:
        tblpr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    _insert_ordered(tblpr, borders)


def _no_borders(table):
    _tbl_borders(table, val="none", sz=0, color="auto")


def _hairlines(table, color=BORDER):
    _tbl_borders(table, val="single", sz=4, color=color)


_EMU_PER_TWIP = 635


def _cant_split(row):
    """Impide que la fila se parta entre dos páginas.

    Un callout partido deja el título solo al pie de una página y el cuerpo en la
    siguiente, que es exactamente lo que un bloque de advertencia no puede hacer.
    """
    _insert_ordered(row._tr.get_or_add_trPr(), OxmlElement("w:cantSplit"))


def _fixed_grid(table, cols_emu: Sequence[int]):
    """Fija el ancho de columna de verdad: layout fijo + ``tblGrid`` reescrito.

    Sin esto Word y LibreOffice reparten el ancho por su cuenta y el ``cell.width``
    que pide python-docx se ignora — así una barra de acento de 0,045" salía ocupando
    media página.
    """
    tblpr = table._tbl.tblPr
    old = tblpr.find(qn("w:tblLayout"))
    if old is not None:
        tblpr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _insert_ordered(tblpr, layout)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), cols_emu):
            gc.set(qn("w:w"), str(int(w / _EMU_PER_TWIP)))


def _p_border(p, edge: str, color: str, size: int, space: int = 1):
    ppr = p._p.get_or_add_pPr()
    borders = _ensure(ppr, "pBdr")
    old = borders.find(qn(f"w:{edge}"))
    if old is not None:
        borders.remove(old)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    # El orden dentro de pBdr también es fijo: top · left · bottom · right · between · bar.
    seq = ["top", "left", "bottom", "right", "between", "bar"]
    rank = seq.index(edge)
    for existing in borders:
        if seq.index(_local(existing.tag)) > rank:
            existing.addprevious(el)
            break
    else:
        borders.append(el)
    return p


def _rule(doc, color=BORDER_STRONG, size=6, space_before=4, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return _p_border(p, "bottom", color, size)


class Doc:
    """Envoltorio de estilo sobre python-docx con la gramática de marca SDQ·MIP."""

    def __init__(self, assets_dir: str):
        self.d = Document()
        self.assets = assets_dir
        self._setup()

    # ── configuración base ──
    def _setup(self):
        st = self.d.styles["Normal"]
        st.font.name = BODY_FONT
        st.font.size = Pt(9.5)
        st.font.color.rgb = rgb(BODY)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.22
        rpr = st.element.get_or_add_rPr()
        rfonts = _ensure(rpr, "rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), BODY_FONT)

        # El template de python-docx trae `<w:zoom w:val="bestFit"/>`, que el esquema
        # rechaza por faltarle `w:percent`. Se corrige acá para que el archivo valide.
        settings = self.d.settings.element
        zoom = settings.find(qn("w:zoom"))
        if zoom is not None and zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")

        for s in self.d.sections:
            s.page_width = Inches(8.5)
            s.page_height = Inches(11)
            s.top_margin = Inches(0.85)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.85)
            s.right_margin = Inches(0.85)

    @property
    def content_width(self) -> int:
        s = self.d.sections[-1]
        return s.page_width - s.left_margin - s.right_margin

    # ── bloques de texto ──
    def eyebrow(self, text: str, color=ACCENT):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        _set_font(p.add_run(text.upper()), font=MONO_FONT, size=7.5, color=color,
                  bold=True, spacing=26)
        return p

    def h1(self, text: str, eyebrow: Optional[str] = None, page_break=True):
        if page_break:
            self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        if eyebrow:
            self.eyebrow(eyebrow)
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        _set_font(p.add_run(text), font=DISPLAY_FONT, size=19, color=INK, bold=True,
                  spacing=-12)
        _rule(self.d, color=ACCENT, size=10, space_before=1, space_after=10)
        return p

    def h2(self, text: str):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        _set_font(p.add_run(text), font=DISPLAY_FONT, size=12.5, color=INK, bold=True,
                  spacing=-8)
        return p

    def h3(self, text: str):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        _set_font(p.add_run(text), font=DISPLAY_FONT, size=10.5, color=ACCENT_INK,
                  bold=True)
        return p

    def p(self, text: str = "", *, size=9.5, color=BODY, italic=False, bold=False,
          space_after=6, space_before=0, align=None):
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(space_after)
        par.paragraph_format.space_before = Pt(space_before)
        if align is not None:
            par.alignment = align
        if text:
            _set_font(par.add_run(text), size=size, color=color, italic=italic, bold=bold)
        return par

    def rich(self, parts: Sequence[Tuple[str, dict]], *, space_after=6, space_before=0):
        """Párrafo con tramos de formato distinto: [(texto, {bold/color/font/size})]."""
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(space_after)
        par.paragraph_format.space_before = Pt(space_before)
        for text, opts in parts:
            _set_font(par.add_run(text), **{"size": 9.5, "color": BODY, **opts})
        return par

    def bullets(self, items: Iterable, *, color=BODY, size=9.5):
        for it in items:
            par = self.d.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.22)
            par.paragraph_format.first_line_indent = Inches(-0.14)
            par.paragraph_format.space_after = Pt(3)
            _set_font(par.add_run("— "), font=BODY_FONT, size=size, color=ACCENT, bold=True)
            if isinstance(it, tuple):
                lead, rest = it
                _set_font(par.add_run(lead), size=size, color=INK, bold=True)
                _set_font(par.add_run(rest), size=size, color=color)
            else:
                _set_font(par.add_run(it), size=size, color=color)

    def callout(self, title: str, text: str, *, tone="accent"):
        fill, bar, ink = {
            "accent": (ACCENT_SOFT, ACCENT, ACCENT_INK),
            "warn": (WARN_SOFT, WARN, WARN),
            "alert": (ALERT_SOFT, ALERT, ALERT),
            "ok": (OK_SOFT, OK, OK),
            "teal": (TEAL_SOFT, TEAL, TEAL),
        }[tone]
        # La barra de acento es una COLUMNA sombreada, no un borde de párrafo: el borde
        # izquierdo de párrafo dentro de una celda queda recortado por el margen y no se
        # dibuja (verificado en la conversión a PDF). Una columna siempre se ve.
        t = self.d.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        _no_borders(t)
        bar_w = Emu(int(Inches(0.05)))
        _fixed_grid(t, [int(bar_w), self.content_width - int(bar_w)])
        _cant_split(t.rows[0])
        left = t.rows[0].cells[0]
        left.width = bar_w
        _cell_shade(left, bar)
        _cell_margins(left, top=0, bottom=0, left=0, right=0)
        left.paragraphs[0].paragraph_format.space_after = Pt(0)

        cell = t.rows[0].cells[1]
        cell.width = Emu(self.content_width - int(bar_w))
        _cell_shade(cell, fill)
        _cell_margins(cell, top=120, bottom=120, left=160, right=160)
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        _set_font(p0.add_run(title.upper()), font=MONO_FONT, size=7.5, color=ink,
                  bold=True, spacing=24)
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        _set_font(p1.add_run(text), size=9.5, color=INK)
        self.p(space_after=4)
        return t

    def code(self, lines: Sequence[str], *, fill=SURFACE2):
        t = self.d.add_table(rows=1, cols=1)
        t.autofit = False
        _no_borders(t)
        _fixed_grid(t, [self.content_width])
        cell = t.rows[0].cells[0]
        cell.width = self.content_width
        _cell_shade(cell, fill)
        _cell_margins(cell, top=110, bottom=110, left=140, right=140)
        first = True
        for line in lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            _set_font(p.add_run(line), font=MONO_FONT, size=7.6, color=INK)
        self.p(space_after=4)
        return t

    # ── tablas ──
    def table(self, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Optional[Sequence[float]] = None, *, size=8.3,
              mono_cols: Sequence[int] = (), header_fill=INK, zebra=True):
        n = len(headers)
        t = self.d.add_table(rows=1, cols=n)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        _hairlines(t)
        total = self.content_width
        if widths:
            scale = sum(widths)
            cols = [int(total * w / scale) for w in widths]
        else:
            cols = [int(total / n)] * n

        _fixed_grid(t, cols)

        hdr = t.rows[0]
        for i, h in enumerate(headers):
            c = hdr.cells[i]
            c.width = Emu(cols[i])
            _cell_shade(c, header_fill)
            _cell_margins(c, top=70, bottom=70)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run(h.upper()), font=MONO_FONT, size=7.0, color="FFFFFF",
                      bold=True, spacing=18)
        _insert_ordered(hdr._tr.get_or_add_trPr(), OxmlElement("w:tblHeader"))

        for ri, row in enumerate(rows):
            tr = t.add_row()
            _cant_split(tr)
            cells = tr.cells
            for i, val in enumerate(row):
                c = cells[i]
                c.width = Emu(cols[i])
                _cell_margins(c, top=60, bottom=60)
                if zebra and ri % 2 == 1:
                    _cell_shade(c, CANVAS)
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                font = MONO_FONT if i in mono_cols else BODY_FONT
                bold = i == 0
                _set_font(p.add_run(str(val)), font=font, size=size,
                          color=INK if bold else BODY, bold=bold)
        self.p(space_after=6)
        return t

    def swatches(self, entries: Sequence[Tuple[str, str, str]], *, cols=3,
                 on_dark=False):
        """entries = [(token, hex, uso)] → grilla de muestras de color."""
        rows = (len(entries) + cols - 1) // cols
        t = self.d.add_table(rows=rows, cols=cols)
        t.autofit = False
        _no_borders(t)
        w = int(self.content_width / cols)
        _fixed_grid(t, [w] * cols)
        for idx, (token, hx, uso) in enumerate(entries):
            r, c = divmod(idx, cols)
            cell = t.rows[r].cells[c]
            cell.width = Emu(w)
            _cell_margins(cell, top=40, bottom=110, left=40, right=90)
            chip = cell.paragraphs[0]
            chip.paragraph_format.space_after = Pt(2)
            _shade(chip._p, hx)
            # Borde fino: sin él, --surface (#FFFFFF) es una muestra invisible.
            edge = BORDER_STRONG if not on_dark else DARK["--border-strong"]
            for side in ("top", "left", "bottom", "right"):
                _p_border(chip, side, edge, 2, space=0)
            _set_font(chip.add_run(" " * 34), font=MONO_FONT, size=11, color=hx)
            p1 = cell.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            _set_font(p1.add_run(token), font=MONO_FONT, size=7.6, color=INK, bold=True)
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(0)
            _set_font(p2.add_run("#" + hx), font=MONO_FONT, size=7.6, color=ACCENT_INK)
            p3 = cell.add_paragraph()
            p3.paragraph_format.space_after = Pt(0)
            _set_font(p3.add_run(uso), size=7.4, color=MUTED)
        # celdas sobrantes
        for idx in range(len(entries), rows * cols):
            r, c = divmod(idx, cols)
            t.rows[r].cells[c].width = Emu(w)
        self.p(space_after=6)
        return t

    def image(self, name: str, width_in: float, caption: Optional[str] = None):
        path = os.path.join(self.assets, name)
        if not os.path.exists(path):
            self.p(f"[asset ausente: {name}]", color=ALERT, italic=True)
            return
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(path, width=Inches(width_in))
        if caption:
            cp = self.d.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            _set_font(cp.add_run(caption), size=7.8, color=MUTED, italic=True)

    # ── encabezado / pie ──
    def running_furniture(self, header_text: str):
        sec = self.d.sections[0]
        sec.different_first_page_header_footer = True

        hp = sec.header.paragraphs[0]
        hp.paragraph_format.space_after = Pt(2)
        _set_font(hp.add_run(header_text), font=MONO_FONT, size=7.0, color=FAINT,
                  spacing=18)
        _rule_in(hp, BORDER)

        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(fp.add_run("SDQ Consulting · SDQ·MIP  ·  "), font=MONO_FONT,
                  size=7.0, color=FAINT)
        _page_field(fp)


def _rule_in(p, color):
    _p_border(p, "bottom", color, 4, space=2)


def _page_field(p):
    for instr, txt in (("PAGE", "1"), (None, " / "), ("NUMPAGES", "1")):
        if instr is None:
            _set_font(p.add_run(txt), font=MONO_FONT, size=7.0, color=FAINT)
            continue
        r = p.add_run()
        _set_font(r, font=MONO_FONT, size=7.0, color=FAINT)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = f" {instr} "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.append(fld_begin)
        r._r.append(instr_el)
        r._r.append(fld_end)


# ─────────────────────────── datos del catálogo ───────────────────────────

def load_catalog() -> Tuple[list, list]:
    """Catálogo y SKUs, leídos del código (nunca de una lista paralela)."""
    from shared.billing.skus import catalog_skus
    from shared.products.registry import PRODUCT_CATALOG
    return list(PRODUCT_CATALOG), list(catalog_skus())


# Índice / sujeto / estado de validación estructural por eje. Los tres primeros campos
# describen el DISEÑO del producto; el cuarto declara qué clase de validación admite el
# eje (ESTADO_BACKTEST del producto), NUNCA el veredicto de la última corrida.
EJES = {
    "banking": ("Perfil SDQ — Ejecución · Resiliencia", "entidad supervisada por la SIB",
                "Motor con corte transversal real (entidad × trimestre) + cohorte de eventos"),
    "macro": ("IRMP — Índice de Riesgo Macro-Político", "país (panel LatAm + Caribe)",
              "Motor de discriminación con corte transversal de países"),
    "monetary_policy": ("Pronóstico de TPM (hold/cut/hike)", "postura del BCRD",
                        "Backtest expanding-window one-step-ahead sobre decisiones de política"),
    "trade": ("Resiliencia comercial · ICE", "flujo de comercio exterior / país",
              "Motor sobre panel de 24 países LatAm+Caribe"),
    "tourism": ("ITT — Índice de Tracción Turística", "país",
                "Sin corte transversal: índice nacional, un sujeto por período"),
    "free_zones": ("IZF — Índice de Zonas Francas", "país",
                   "Sin corte transversal: agregado nacional CNZFE"),
    "energy": ("IRSE — Índice de Resiliencia del Sistema Eléctrico", "país",
               "Sin corte transversal: falta serie por distribuidora o circuito"),
    "telecom": ("IDT — Índice de Desarrollo de Telecom", "país",
                "Dato pendiente: el boletín INDOTEL está congelado en 2022-Q1"),
    "construction": ("ICC — Índice de Coyuntura de la Construcción", "país",
                     "Sin ejecución realizada de permisos ni corte provincial con historia"),
    "agribusiness": ("IAI · SGPS — atractivo y momentum sectorial", "rama económica",
                     "Motor sobre ramas × años; se ofrece como DESCRIPTIVO"),
    "esg": ("IRC — Índice de Resiliencia Climática", "país",
            "Motor transversal de países contra mortalidad por desastres"),
    "pension": ("ISA — Índice de Solidez de la AFP", "AFP",
                "Motor transversal de AFP × períodos"),
    "insurance": ("ISF — Índice de Solidez Financiera", "aseguradora",
                  "Motor transversal de aseguradoras × años auditados"),
    "economic_structure": ("Estructura del PIB por sectores de origen", "sector económico",
                           "Descriptivo: no hay índice sintético que validar"),
    "social_dev": ("IDM — Índice de Desarrollo Multidimensional",
                   "región de desarrollo / provincia",
                   "Validez CONVERGENTE contra el IDH regional del PNUD"),
    "law": ("Cobertura y cumplimiento de metas del instrumento", "instrumento normativo",
            "Autorreferencial: el desenlace de una ley es lo que el eje ya mide"),
}


# Orden de presentación del estado de validación (§7.3), con el rótulo corto del eje.
_EJE_NOMBRES = [
    ("banking", "Banking Intelligence"),
    ("macro", "Macro & Country Risk"),
    ("monetary_policy", "Política Monetaria"),
    ("trade", "Trade & Logistics"),
    ("esg", "ESG & Climate"),
    ("pension", "Pensiones (SIPEN)"),
    ("insurance", "Seguros (SIS)"),
    ("social_dev", "Desarrollo Social"),
    ("agribusiness", "Agribusiness / sectorial"),
    ("tourism", "Tourism Intelligence"),
    ("free_zones", "Free Zones & Manufacturing"),
    ("energy", "Energy Intelligence"),
    ("telecom", "Telecom Intelligence"),
    ("construction", "Construction Intelligence"),
    ("economic_structure", "Estructura de la Economía"),
    ("law", "Evaluación de Leyes"),
]


PRECIOS = [
    ("Insight · un sector", "insight:{sector}", "mensual", "US$149",
     "Suscripción por eje; entidad nombrada."),
    ("Insight · un sector", "insight:{sector}", "anual", "US$1,490",
     "Equivale a 10 meses; ~17% de descuento."),
    ("Deep Dive · un sector", "deep_dive:{sector}", "una vez", "US$450",
     "Informe on-demand, 8–15 páginas."),
    ("All-Access", "all_access", "mensual", "US$690",
     "Insight de los 16 ejes."),
    ("All-Access", "all_access", "anual", "US$6,900", "Bundle anual."),
    ("Enterprise", "enterprise", "mensual", "US$1,450",
     "Catálogo completo: Insight + Deep Dive."),
    ("Enterprise", "enterprise", "anual", "US$14,500", "Bundle anual."),
    ("Research a Medida", "special:research-custom", "por encargo", "US$3,500",
     "ANCLA decidida (provisional). Consultoría tradicional: US$7,000–25,000+."),
]


# ─────────────────────────── contenido ───────────────────────────

def cover(doc: Doc, hoy: str):
    d = doc.d
    d.add_paragraph().paragraph_format.space_after = Pt(24)
    doc.image("lockup_light.png", 4.4)
    d.add_paragraph().paragraph_format.space_after = Pt(18)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run("DOSSIER DE PRODUCTO Y MARCA"), font=MONO_FONT, size=8.5,
              color=ACCENT, bold=True, spacing=40)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _set_font(p.add_run("Handoff para el diseño de la\npresentación comercial"),
              font=DISPLAY_FONT, size=30, color=INK, bold=True, spacing=-24)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    _set_font(p.add_run("Todo lo que la plataforma hace hoy, con qué se puede afirmar, "
                        "cómo se cobra y con qué reglas visuales se presenta."),
              size=11, color=MUTED)

    _rule(d, color=ACCENT, size=12, space_before=0, space_after=14)

    t = d.add_table(rows=1, cols=4)
    t.autofit = False
    _no_borders(t)
    w = int(doc.content_width / 4)
    _fixed_grid(t, [w] * 4)
    meta = [("Destinatario", "Diseñador senior"),
            ("Emisor", "SDQ Consulting"),
            ("Fecha", hoy),
            ("Estado", "Vigente al corte")]
    for i, (k, v) in enumerate(meta):
        c = t.rows[0].cells[i]
        c.width = Emu(w)
        _cell_margins(c, left=0, right=60)
        p0 = c.paragraphs[0]
        p0.paragraph_format.space_after = Pt(1)
        _set_font(p0.add_run(k.upper()), font=MONO_FONT, size=7, color=FAINT,
                  bold=True, spacing=20)
        p1 = c.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        _set_font(p1.add_run(v), font=DISPLAY_FONT, size=10.5, color=INK, bold=True)

    d.add_paragraph().paragraph_format.space_after = Pt(16)
    doc.callout(
        "Lea esto primero",
        "Este documento describe una plataforma que ya existe y corre. La regla que lo "
        "gobierna —y que la presentación debe heredar— es que SDQ·MIP no afirma más de lo "
        "que su código sostiene. Hay una sección entera (§7) dedicada a qué se puede decir "
        "y qué no; ninguna pieza de venta debería salir sin pasar por ella.",
        tone="accent")


def indice(doc: Doc):
    doc.h1("Contenido", eyebrow="Índice")
    # Índice MANUAL, no un campo TOC de Word: un campo se rellena al abrir el archivo y
    # sale vacío en el PDF, que es la mitad de la entrega.
    doc.table(
        ["§", "Sección", "Qué resuelve"],
        [("0", "Cómo usar este documento", "Las tres capas del handoff y lo que hay que "
                                           "saber antes de abrir el archivo de diseño."),
         ("1", "Qué es SDQ·MIP", "Posicionamiento, el foso, audiencias y la arquitectura "
                                 "en una lámina."),
         ("2", "La marca", "Identidad Arco, tokens de color claro y oscuro, tipografía, "
                           "componentes, visualización, reglas duras y tono."),
         ("3", "El catálogo — 16 ejes", "Los productos, con índice, sujeto, fuente y "
                                        "ángulo de venta."),
         ("4", "Niveles y anatomía del informe", "Pulse · Insight · Deep Dive y las diez "
                                                 "secciones del reporte."),
         ("5", "Capacidades transversales", "Narrativa con guardas, alertas, Data API, "
                                            "research, herramientas, gate de publicación, "
                                            "acceso y facturación."),
         ("6", "Modelo comercial", "SKU, tarifario propuesto, estado del cobro y el "
                                   "argumento de la escalera."),
         ("7", "Qué se puede afirmar y qué no", "El filtro obligatorio de mensaje. "
                                                "Léalo antes de escribir una sola lámina."),
         ("8", "Brief para el diseñador", "Encargo, guion lámina por lámina, aciertos y "
                                          "errores, assets y decisiones pendientes."),
         ("9", "Anexos", "Tokens completos, vectorial del símbolo, navegación, glosario "
                         "y fuentes.")],
        widths=[0.3, 1.7, 3.8], size=8.6)


def seccion_0(doc: Doc):
    doc.h1("Cómo usar este documento", eyebrow="§0 · Instrucciones de handoff")
    doc.p("El documento tiene tres capas y cada una responde a una pregunta distinta del "
          "trabajo de diseño.")
    doc.bullets([
        ("Capas 1–2 (§1–§2) · La marca. ",
         "Identidad, tokens de color claro y oscuro, tipografía, componentes y reglas duras. "
         "Es lo que se aplica; no es una propuesta abierta."),
        ("Capas 3–6 (§3–§6) · El producto y el negocio. ",
         "Los 16 ejes del catálogo, las capacidades transversales, los tres niveles "
         "comerciales y el modelo de precios. Es la materia de las láminas."),
        ("Capa 7 (§7) · Las restricciones de mensaje. ",
         "Qué afirmación está permitida por producto. Es un filtro obligatorio, no una "
         "sugerencia de tono."),
    ])
    doc.p()
    doc.h2("Tres cosas que hay que saber antes de abrir el archivo de diseño")
    doc.bullets([
        ("La marca visual ya está construida y corre en producción. ",
         "Los tokens de §2.2 son literalmente los que consume la aplicación. Reprodúzcalos "
         "con exactitud: el objetivo no es reinventar la paleta, es que la presentación y "
         "el producto se lean como la misma cosa."),
        ("Hay un solo sistema de color, y es el de §2.2. ",
         "Hasta hace poco la aplicación y los informes tenían paletas distintas; se "
         "unificaron, y ninguna superficie declara ya un color propio. §2.4 cuenta qué "
         "cambió, por si aparece material viejo."),
        ("Las cifras de validación NO se copian de este documento. ",
         "No hay ninguna: son la única clase de dato que se lee en vivo de la plataforma. "
         "§7.4 dice exactamente de dónde y por qué."),
    ])
    doc.p()
    doc.callout(
        "Convención de nombre",
        "La marca se escribe «SDQ·MIP» con punto medio (U+00B7), nunca «SDQMIP» ni "
        "«SDQ MIP». En el logotipo, «SDQ» va en tinta y «·MIP» en gris apagado. El nombre "
        "extendido es «SDQ Market Intelligence Platform»; el emisor es SDQ Consulting.",
        tone="teal")


def seccion_1(doc: Doc):
    doc.h1("Qué es SDQ·MIP", eyebrow="§1 · Posicionamiento")

    doc.callout(
        "La frase de una sola línea",
        "SDQ·MIP es la plataforma de inteligencia de mercado de la República Dominicana y "
        "el Caribe: dieciséis ejes de análisis construidos sobre la fuente oficial de cada "
        "sector, donde toda cifra publicada declara si es dato medido, criterio de la casa "
        "o brecha — y se puede auditar hasta su origen.",
        tone="accent")

    doc.h2("1.1 El problema de mercado")
    doc.p("En República Dominicana el dato existe —la SIB, el BCRD, la SIPEN, la "
          "Superintendencia de Seguros, la DGA, la ONE y el MEPyD publican— pero está "
          "fragmentado en PDF, Excel y boletines con criterios distintos, ninguna serie "
          "normalizada y ninguna lectura comparable entre períodos. Quien necesita decidir "
          "—un comité de crédito, un fondo, un gremio, un multilateral, un regulador— "
          "paga a un consultor para que reconstruya a mano lo que ya es público, y recibe "
          "un informe que no se puede auditar y que envejece el día que se entrega.")

    doc.h2("1.2 El foso: procedencia auditable")
    doc.p("El diferenciador de SDQ·MIP no es tener el dato. Es tener el dato con su "
          "procedencia declarada a nivel de variable, y una plataforma que se niega a "
          "publicar lo que no puede sostener. Cuatro reglas lo hacen operativo:")
    doc.bullets([
        ("Se declara la brecha, nunca se rellena. ",
         "Un dato ausente es ausencia, jamás un cero ni un promedio. Y si una métrica existe "
         "pero no mide lo que el eje afirma medir para esa entidad, tampoco se publica: se "
         "declara el motivo."),
        ("Las relaciones se computan, no se narran. ",
         "Dirección, superlativos, deltas, rankings y posiciones se calculan en código; el "
         "modelo de lenguaje los copia. Un modelo acierta las cifras y falla las relaciones."),
        ("Solo se ordena lo comparable. ",
         "Un score armado sobre 3 de 5 dimensiones no rankea contra uno de 5. Los parciales "
         "no se ocultan —eso los haría desaparecer sin aviso—: van aparte y marcados."),
        ("La frescura veta. ",
         "Una credencial de validación cuyo insumo cambió después del cálculo no se publica; "
         "y una cuya frescura es indeterminada, tampoco. «No sé de cuándo es» y «está al "
         "día» son cosas distintas."),
    ])
    doc.p()
    doc.rich([
        ("Por qué esto vende. ", {"bold": True, "color": INK}),
        ("Un comprador institucional puede auditar. El mayor riesgo reputacional de una "
         "casa de inteligencia es la brecha entre lo que el material promete y lo que el "
         "producto sostiene. SDQ·MIP convierte esa exposición en el argumento: si el "
         "comprador pregunta «¿esto es dato o criterio?», la respuesta existe a nivel de "
         "variable y está rotulada en la interfaz, en el informe y en la API.", {}),
    ])

    doc.h2("1.3 A quién se le vende")
    doc.table(
        ["Audiencia", "Qué compra", "Nivel típico"],
        [
            ("Bancos y entidades financieras",
             "Posición propia y de pares en el panel supervisado; alerta temprana; "
             "insumo para comité de riesgo.", "Insight · Enterprise"),
            ("Fondos, inversionistas y family offices",
             "Riesgo país, atractivo sectorial, resiliencia comercial y climática; "
             "due diligence de entrada.", "Insight · Deep Dive"),
            ("Gremios y asociaciones sectoriales",
             "Pulso del sistema sin nombrar entidades; posicionamiento del sector ante "
             "el Estado y la prensa.", "Pulse · Insight"),
            ("Multilaterales y cooperación",
             "Desarrollo social sub-nacional, ESG y clima, evaluación de cumplimiento de "
             "instrumentos normativos.", "Deep Dive · Research"),
            ("Estado, reguladores y comisiones legislativas",
             "Evaluación de leyes contra las metas que ellas mismas fijaron; estructura "
             "de la economía.", "Deep Dive · Research"),
            ("Empresas y corporativos",
             "Contexto de marca y categoría, escenarios, Deal Scoring de una operación "
             "concreta.", "Insight · herramientas"),
            ("Prensa económica y academia",
             "Pulse abierto: bandas y tendencias del sistema, sin entidades nombradas.",
             "Pulse"),
        ],
        widths=[1.5, 3.2, 1.1])

    doc.h2("1.4 La arquitectura, en una lámina")
    doc.image("arch.png", 6.8,
              "Figura 1 — De la fuente oficial al artefacto vendible. Cada capa es "
              "auditable hacia atrás: toda cifra publicada rastrea hasta su conector, "
              "su peso de doctrina y su estado de validación.")


def seccion_2(doc: Doc):
    doc.h1("La marca", eyebrow="§2 · Sistema de identidad y diseño")
    doc.p("La dirección visual aprobada se llama «Claro & Vivo» y la identidad, «Arco». "
          "Ambas están implementadas en la aplicación: lo que sigue no es una propuesta, "
          "es el contrato vigente.")

    # ── 2.1 identidad
    doc.h2("2.1 El símbolo: Arco")
    doc.p("El logotipo es un arco abierto —la aguja de un medidor 0–100— con un punto: "
          "el dato, la señal. Resume el producto, porque un medidor explicable es "
          "literalmente lo que la plataforma emite. En la aplicación el arco puede "
          "animarse como medidor.")
    doc.image("logo_arco_1024.png", 1.25)
    doc.p("Hasta hace poco circulaban tres marcas distintas. Ya no: quedó una, y el "
          "diseñador debería saber cuáles se retiraron para reconocer material viejo.",
          space_after=4)
    doc.image("logo_variants.png", 6.4,
              "Figura 2 — Izquierda, las dos retiradas: el SVG que servía de favicon "
              "dibujaba el arco con un círculo punteado y el hueco quedaba tan corto que "
              "el punto se fundía contra el terminal; el logotipo de los informes tenía la "
              "construcción correcta pero en otro azul. Derecha, la canónica: la "
              "construcción del de informes en el azul acento de la aplicación.")
    doc.callout(
        "Resuelto — y protegido por un test",
        "El símbolo quedó en UNA geometría y UN azul (#1E6FFF). La geometría vive en un "
        "solo lugar del código y las copias inevitables —el componente de la aplicación, "
        "el favicon, el PNG de los informes— se comparan contra ella en cada corrida de "
        "integración continua: si una deriva, el build falla. Si encuentra material con el "
        "azul #2B6CB0 o con el punto pegado al arco, está desactualizado.",
        tone="ok")
    doc.p()
    doc.rich([("Nota de origen. ", {"bold": True, "color": INK}),
              ("SDQ no tiene un logotipo corporativo oficial. «Arco» es una propuesta de "
               "producto —hay cuatro direcciones exploradas (Arco · Mira · Índice · "
               "Cuadrante) y Arco es la recomendación—. La presentación puede tratarla "
               "como identidad del producto SDQ·MIP, no como identidad de SDQ "
               "Consulting.", {})])

    doc.h3("Logotipo completo (símbolo + palabra)")
    doc.image("lockup_light.png", 3.9)
    doc.image("lockup_dark.png", 3.9,
              "Figura 3 — Bloque de marca sobre fondo claro y sobre fondo oscuro. "
              "«SDQ» en tinta, «·MIP» en apagado. Bajada opcional en monoespaciada, "
              "mayúsculas, con interletrado abierto.")
    doc.bullets([
        ("Área de resguardo: ", "una altura de símbolo por los cuatro lados."),
        ("Tamaño mínimo: ", "símbolo 20 px / 6 mm; bloque completo 90 px / 25 mm."),
        ("Sobre fotografía o color: ", "solo el símbolo en blanco sólido, sin caja."),
        ("Prohibido: ", "estirar, rotar, cambiar el azul del contenedor, poner el arco "
                        "en degradado o agregarle sombra."),
    ])

    # ── 2.2 color claro
    doc.h2("2.2 Color — tema claro (canónico)")
    doc.p("Toda la interfaz consume variables CSS; nunca un hex suelto. La presentación "
          "debe usar exactamente estos valores. El tema claro es el canónico de la marca: "
          "superficies claras, tinta navy solo para texto, azul eléctrico como acento. "
          "El color comunica estado; no decora.")
    doc.h3("Superficies y tinta")
    doc.swatches([
        ("--canvas", CANVAS, "Fondo de página"),
        ("--surface", SURFACE, "Tarjetas y paneles"),
        ("--surface-2", SURFACE2, "Fondo hundido, código"),
        ("--ink", INK, "Títulos y cifras"),
        ("--body", BODY, "Texto corrido"),
        ("--muted", MUTED, "Subtítulos, leyendas"),
        ("--faint", FAINT, "Metadato terciario"),
        ("--border", BORDER, "Bordes de tarjeta"),
        ("--border-strong", BORDER_STRONG, "Campos, divisores"),
    ])
    doc.h3("Acento y estado")
    doc.swatches([
        ("--accent", ACCENT, "Primario, foco, activo"),
        ("--accent-hover", ACCENT_HOVER, "Estado hover"),
        ("--accent-soft", ACCENT_SOFT, "Fondo de acento"),
        ("--accent-ink", ACCENT_INK, "Texto sobre acento suave"),
        ("--teal", TEAL, "Acento secundario"),
        ("--teal-soft", TEAL_SOFT, "Fondo teal"),
        ("--ok", OK, "Positivo / fuerte"),
        ("--warn", WARN, "Vigilar"),
        ("--alert", ALERT, "Crítico / débil"),
    ])
    doc.h3("Paleta de datos (series de gráfico)")
    doc.p("Máximo seis series; después se reutilizan en orden. Nunca se colorea por "
          "gusto: el orden es fijo para que dos gráficos distintos del mismo informe "
          "usen el mismo color para la primera serie.", space_after=4)
    doc.swatches([
        ("--c1", C1, "Serie 1"), ("--c2", C2, "Serie 2"), ("--c3", C3, "Serie 3"),
        ("--c4", C4, "Serie 4"), ("--c5", C5, "Serie 5"), ("--c6", C6, "Serie 6"),
    ], cols=6)

    # ── 2.3 color oscuro
    doc.h2("2.3 Color — tema oscuro (ciudadano de primera clase)")
    doc.p("El tema oscuro no es una cortesía: es un modo real, persistido por usuario, y "
          "la presentación debería mostrarlo. Todo componente y todo gráfico se resuelven "
          "por token, así que cambiar de tema no toca el marcado.")
    doc.swatches([
        ("--canvas", DARK["--canvas"], "Fondo de página"),
        ("--surface", DARK["--surface"], "Tarjetas"),
        ("--surface-2", DARK["--surface-2"], "Fondo hundido"),
        ("--ink", DARK["--ink"], "Títulos y cifras"),
        ("--body", DARK["--body"], "Texto corrido"),
        ("--muted", DARK["--muted"], "Subtítulos"),
        ("--border", DARK["--border"], "Bordes"),
        ("--accent", DARK["--accent"], "Primario"),
        ("--accent-ink", DARK["--accent-ink"], "Texto sobre acento"),
        ("--ok", DARK["--ok"], "Positivo"),
        ("--warn", DARK["--warn"], "Vigilar"),
        ("--alert", DARK["--alert"], "Crítico"),
    ], cols=4)
    doc.p("Series de datos en oscuro: ", space_after=2)
    doc.swatches([
        ("--c1", DARK["--c1"], ""), ("--c2", DARK["--c2"], ""), ("--c3", DARK["--c3"], ""),
        ("--c4", DARK["--c4"], ""), ("--c5", DARK["--c5"], ""), ("--c6", DARK["--c6"], ""),
    ], cols=6)

    # ── 2.4 divergencia
    doc.h2("2.4 Una sola paleta en pantalla y en papel")
    doc.p("Durante un tiempo hubo dos marcas: la aplicación en «Claro & Vivo» y los "
          "informes en un navy con acento rojo que la aplicación ya no usaba. El cliente "
          "veía un producto azul eléctrico en pantalla y recibía un PDF en otra marca. "
          "Ya está unificado, y la unificación es estructural: ninguna superficie declara "
          "un color — todas leen los mismos tokens.")
    doc.table(
        ["Superficie", "Antes", "Hoy"],
        [("Aplicación web", "#0A1A3A · #1E6FFF", "sin cambios — era la referencia"),
         ("Informe PDF", "#1A365D · #2B6CB0 · #E11D48", "#0A1A3A · #1E6FFF"),
         ("Informe Word", "#1A365D · #2B6CB0", "#0A1A3A · #1E6FFF"),
         ("PDF forense de banca", "hex propios, escritos a mano", "#0A1A3A · #1E6FFF"),
         ("Informe de Contexto de Marca", "#0B1F3A · #D7263D", "#0A1A3A · #1E6FFF"),
         ("Comprobante fiscal", "grises propios", "tokens de marca"),
         ("Logotipo de los informes", "#2B6CB0", "#1E6FFF")],
        widths=[1.7, 1.9, 1.9])
    doc.callout(
        "Qué pasó con el rojo",
        "El «signal red» #E11D48 se RETIRÓ como acento decorativo: donde marcaba una cita "
        "destacada, ahora va el azul de marca. El rojo sobrevive únicamente como color "
        "SEMÁNTICO —#C8392E, el token de alerta— y solo con significado: un valor negativo "
        "en un gráfico, la estampa de «MUESTRA», la leyenda de exención de ITBIS en un "
        "comprobante. Si en la presentación aparece rojo, tiene que estar diciendo algo.",
        tone="alert")
    doc.p()
    doc.rich([
        ("Un detalle de accesibilidad que conviene respetar. ", {"bold": True, "color": INK}),
        ("El acento puro #1E6FFF sobre blanco da 4,40:1 de contraste, por debajo del "
         "4,5:1 que pide WCAG AA para texto normal. Por eso el acento pinta RELLENOS "
         "—barras, filetes, fondos— y el texto en acento usa #1551C0 (7,07:1). No es una "
         "sutileza de implementación: aplica igual en las láminas.", {}),
    ])

    # ── 2.5 tipografía
    doc.h2("2.5 Tipografía")
    doc.table(
        ["Rol", "Familia", "Peso / tamaño", "Detalle"],
        [("Display y títulos", "Plus Jakarta Sans", "700 / 800", "Interletrado −0.02em"),
         ("Cuerpo e interfaz", "Inter", "400 / 500 / 600", "Interlínea 1.5"),
         ("Cifras, códigos, etiquetas", "JetBrains Mono", "400 / 500 / 600",
          "tabular-nums obligatorio"),
         ("Antetítulo (kicker)", "JetBrains Mono", "600 · 10–11 px",
          "Mayúsculas, interletrado .12–.2em, en acento"),
         ("Métrica grande", "Plus Jakarta Sans o Mono", "700 / 800 · 24–34 px",
          "Siempre tabular"),
         ("Título de página", "Plus Jakarta Sans", "800 · 26 px", "Interlínea cerrada"),
         ("Título de tarjeta", "Plus Jakarta Sans", "700 · 15 px", "Una sola línea"),
         ("Cuerpo", "Inter", "400–500 · 13–14 px", "—"),
         ("Leyenda", "Inter", "400 · 12 px", "En --muted")],
        widths=[1.4, 1.5, 1.5, 1.7])
    doc.bullets([
        ("Mínimos: ", "nunca por debajo de 12 px en interfaz; las cifras clave nunca por "
                      "debajo de 24 px."),
        ("Cifras tabulares en todo dato numérico. ",
         "Una columna de números que baila al cambiar de fila delata amateurismo en un "
         "producto financiero."),
        ("Las tres familias son de Google Fonts y de licencia abierta ",
         "(SIL Open Font License), así que se pueden incrustar en el archivo de "
         "presentación sin restricción."),
    ])

    # ── 2.6 componentes
    doc.h2("2.6 Componentes y gramática de pantalla")
    doc.image("ui_light.png", 6.7)
    doc.image("ui_dark.png", 6.7,
              "Figura 4 — Muestra de componentes en ambos temas: medidor de arco, tarjeta "
              "de métrica con delta y sellos de procedencia, vocabulario de bandas, "
              "botones y tabla de ranking con cifras tabulares.")
    doc.table(
        ["Componente", "Definición"],
        [("Card", "Superficie, borde de 1 px, radio 16 px, sombra suave de tarjeta."),
         ("CardHead", "Icono de 30 px sobre acento suave + título en UNA línea + subtítulo."),
         ("PageHead", "Antetítulo + título + bajada de ancho máximo + ranura de acciones."),
         ("Button", "Tres variantes: primario (acento), fantasma (borde), suave "
                    "(acento suave). Alto 36–38 px, radio 10 px, peso 600."),
         ("StatTile", "Etiqueta + métrica grande tabular + delta con flecha y color de signo."),
         ("Gauge", "Arco 0–100 en SVG propio, color por banda. Es el motivo del logotipo."),
         ("BandBadge", "Píldora con la banda del índice; color por banda."),
         ("Delta", "Signo + flecha; verde sube, rojo baja, gris neutro. Tabular."),
         ("Chip / Tag", "Píldora de 12 px con borde suave; variante con punto de estado."),
         ("Tabs", "Subrayado activo en acento."),
         ("Segmented", "Control de 2–3 opciones (p. ej. Base 100 / Valor)."),
         ("StateBlock", "Los cuatro estados obligatorios de pantalla."),
         ("Skeleton", "Destello sobre superficie hundida durante la carga."),
         ("Toast", "Inferior centrado, fondo tinta.")],
        widths=[1.2, 4.6])

    doc.h3("Bandas de índice — el vocabulario de estado")
    doc.p("Los índices de 0 a 100 se leen por banda, y el vocabulario cambia según lo que "
          "el eje mide. Nunca se mezclan dos escalas en la misma lectura.", space_after=4)
    doc.table(
        ["Escala", "Bandas, de mejor a peor", "Naturaleza"],
        [("Índice general", "Fuerte ≥85 · Sólido 70–84 · Vigilar 55–69 · Débil <55",
          "Absoluta"),
         ("Ejes de riesgo (IRMP, IRC)",
          "Riesgo bajo ≥80 · Riesgo moderado 60–79 · Riesgo elevado 40–59 · Riesgo alto <40",
          "Invertida: mayor score = MENOR riesgo. Ojo, los cortes NO son los de la escala "
          "general — no se reusan"),
         ("Resiliencia (financiero)",
          "Sólida ≥75 · Adecuada 60–74 · En vigilancia 45–59 · Frágil <45",
          "ABSOLUTA — cortes fijos, iguales en los cuatro sectores financieros"),
         ("Ejecución (financiero)",
          "Sobresaliente · Competitiva · Rezagada · Deficiente",
          "RELATIVA al panel comparable — cuartiles por tipo de entidad, expuestos para "
          "ser auditables"),
         ("Índices sectoriales (IAI, SGPS)", "La banda general, sobre el score 0–100",
          "No hay grado por letra: la escala A–D nunca llegó a producción"),
         ("Cualquiera, sin dato", "Sin dato", "La brecha se declara en la insignia; nunca "
          "se rellena con un cero ni se esconde")],
        widths=[1.15, 2.5, 2.15])
    doc.callout(
        "Cambio de marca importante y reciente",
        "La notación por letras SDQ-AAA … SDQ-D fue RETIRADA y no se publica más. Usaba la "
        "gramática de una calificadora de riesgo regulada sin serlo, y «SDQ-D» —la etiqueta "
        "más grave del vocabulario de agencia— cubría 45 puntos de rango, alcanzando a "
        "entidades que operan con normalidad. La reemplaza el «Perfil SDQ»: dos ejes "
        "independientes, Ejecución y Resiliencia, cada uno de 0 a 100 con su propia banda, "
        "que NO se resumen en un símbolo único. Si encuentra material con SDQ-AA+ o "
        "similar, está desactualizado.",
        tone="alert")
    doc.p()
    doc.p("Cuatro reglas al maquetar el Perfil SDQ:", space_after=3)
    doc.bullets([
        ("Nunca se resumen en un símbolo ni en un número. ",
         "Es lo que el sistema de dos ejes existe para evitar: una entidad puede ser "
         "sólida y poco eficiente a la vez, y el lector tiene que ver las dos cosas."),
        ("Van juntos y con el mismo peso visual. ",
         "Jerarquizar uno reintroduce el problema por diseño."),
        ("Con universo chico, la banda sola engaña. ",
         "«Sobresaliente» entre 4 dice bastante menos que entre 42: por debajo de 15 "
         "entidades la interfaz muestra la posición relativa al lado."),
        ("Un movimiento se lee dentro de su propio eje. ",
         "«Sobresaliente» y «Sólida» son de escalas distintas y no se comparan."),
    ])

    # ── 2.7 dataviz
    doc.h2("2.7 Visualización de datos")
    doc.bullets([
        ("Densidad de terminal financiera, ", "no de tablero genérico. La referencia es "
         "Koyfin, Linear y Trading Economics: jerarquía tipográfica fuerte, cifras "
         "tabulares, cromo mínimo."),
        ("Piezas estándar: ", "área y línea, barras, multi-serie base 100, dispersión de "
         "burbujas, radar."),
        ("Piezas a medida (SVG propio): ", "medidor de arco, mapa de calor "
         "indicador×trimestre, treemap, cartograma de mosaicos, distribución por tramos, "
         "abanico de escenarios, barras de driver ±, medidores de sub-componente, barras "
         "de ranking, sparkline."),
        ("Sin geografía falsa. ", "Nunca un mapa geográfico real: cartograma de mosaicos "
         "etiquetados. Un mapa sugiere una precisión territorial que el dato no tiene."),
        ("Ejes y tooltips sobrios, ", "rejilla tenue, sin sombras pesadas."),
    ])

    # ── 2.8 reglas
    doc.h2("2.8 Reglas duras (no negociables)")
    doc.table(
        ["#", "Regla", "Por qué"],
        [("1", "Cabeceras a una sola línea, con recorte por elipsis.",
          "Un título que envuelve junto a una insignia hermana se solapa con el subtítulo "
          "en el borde exacto del salto de línea. Ya ocurrió."),
         ("2", "Si el nombre es dinámico y largo (entidad, región, sector), va en el "
               "subtítulo, no en el título.", "El subtítulo sí puede envolver."),
         ("3", "Cuatro estados por pantalla, siempre: cargando, vacío, error, sin permiso.",
          "Una pantalla con un solo estado es una pantalla que miente cuando falla."),
         ("4", "Cifras tabulares en todo dato numérico.", "Producto financiero."),
         ("5", "Todo color por token; ningún hex suelto.", "Es lo que hace real el tema oscuro."),
         ("6", "Gráficos conscientes del tema.", "Idem."),
         ("7", "Áreas táctiles de 44 px o más en móvil.", "Accesibilidad."),
         ("8", "Tema, ruta, período y ámbito persistidos.", "Continuidad de sesión."),
         ("9", "Interfaz en español; identificadores en inglés; narrativa en ES/EN/FR.",
          "Convención del repositorio."),
         ("10", "Sin emoji. Sin degradados decorativos. Sin tarjetas con borde de acento "
                "a la izquierda.", "Mínimo viable de cromo.")],
        widths=[0.3, 2.8, 2.7])

    doc.h2("2.9 Tono y voz")
    doc.bullets([
        ("Institucional, preciso, sobrio. ", "Sin marketing vacío ni adjetivos de relleno."),
        ("Investigativo, no descriptivo. ", "El material concluye; explica el «y por "
         "tanto», no reexpone cifras."),
        ("Cada elemento gana su lugar. ", "Nada de estadísticas o iconos de relleno."),
        ("La honestidad de procedencia es visible y es el argumento, ",
         "no una disculpa al pie."),
        ("Marco narrativo SCQA ", "(Situación · Complicación · Pregunta · Respuesta) para "
         "las lecturas de analista."),
    ])


def seccion_3(doc: Doc, catalogo):
    doc.h1("El catálogo — 16 ejes", eyebrow="§3 · Productos")
    doc.p("Cada eje es un producto independiente con su propia fuente autoritativa, su "
          "motor de índice y sus tres niveles comerciales. Ninguno importa de otro: se "
          "comunican por eventos. El catálogo se lee del código, no de una lista paralela.")

    filas = []
    for e in catalogo:
        indice, sujeto, _ = EJES.get(e.sector_key, ("—", "—", "—"))
        filas.append((e.display_name.replace("SDQ ", ""), indice, sujeto, e.source))
    doc.table(["Producto", "Índice o lectura", "Sujeto", "Fuente autoritativa"], filas,
              widths=[1.35, 1.75, 1.25, 1.55], size=7.6)

    doc.callout(
        "Cómo dimensionar el catálogo en la presentación",
        "Son DIECISÉIS ejes. Cualquier material que diga «catorce» o «diez» quedó "
        "desactualizado: el catálogo creció con Desarrollo Social y Evaluación de Leyes. "
        "La fuente canónica es el registro de productos del código, y por eso este "
        "documento la lee de ahí en vez de transcribirla.",
        tone="teal")

    doc.h2("3.1 Qué hace único a cada eje")
    doc.p("Notas de venta por producto — el ángulo que lo distingue de un informe "
          "genérico de sector.")
    doc.table(
        ["Eje", "El ángulo de venta"],
        [
            ("Banking Intelligence",
             "Cubre TODO el universo supervisado por la SIB, no solo bancos múltiples: "
             "también ahorro y crédito, corporaciones, asociaciones, intermediación "
             "cambiaria y fiduciarias. 19–21 indicadores en cinco sub-componentes, con "
             "recalibración de pesos por tipo de entidad. Perfil SDQ de dos ejes, "
             "alerta temprana anclada a la crisis de 2003, y un módulo aparte para "
             "fideicomisos con su propio índice de salud."),
            ("Macro & Country Risk",
             "Índice de riesgo país sobre dato real de gobernanza (WGI), macro (WDI), "
             "calificación soberana y tensión de eventos (GDELT), más tres variables "
             "institucionales de juicio experto declarado. Panel LatAm y Caribe para "
             "comparar, no solo RD."),
            ("Política Monetaria",
             "El único producto con modelo predictivo de verdad: regla tipo Taylor "
             "interpretable más clasificador XGBoost hold/cut/hike, con backtest "
             "expanding-window one-step-ahead y bitácora de aciertos en vivo."),
            ("Trade & Logistics",
             "Aduanas (DGA) normalizadas con socios comerciales y resiliencia del flujo; "
             "el motor se valida sobre 24 países del panel regional, que es donde hay "
             "poder estadístico."),
            ("Tourism Intelligence",
             "Índice de tracción turística sobre llegadas, gasto y capacidad; el sector "
             "de mayor peso en la conversación económica del país."),
            ("Free Zones & Manufacturing",
             "Dinamismo exportador, inversión, empleo y productividad de zonas francas "
             "sobre el dato de CNZFE — el motor exportador que ningún tablero público "
             "sigue en serie."),
            ("Energy Intelligence",
             "Resiliencia del sistema eléctrico: matriz de generación, calidad de "
             "servicio y dependencia de combustible, sobre SIE y el Organismo Coordinador."),
            ("Telecom Intelligence",
             "Desarrollo del sector sobre INDOTEL e ITU. Nota de honestidad: el boletín "
             "de INDOTEL está congelado desde 2022-Q1 y la plataforma lo declara en vez "
             "de disimularlo."),
            ("Construction Intelligence",
             "Coyuntura de la construcción sobre permisos de MIVHED y cuentas del BCRD; "
             "serie corta (2022 en adelante) y declarada como tal."),
            ("Agribusiness",
             "Atractivo de inversión (IAI) y momentum (SGPS) de la rama agropecuaria, "
             "dentro del corte transversal sectorial del BCRD."),
            ("ESG & Climate",
             "Resiliencia climática nacional con panel Caribe/LatAm: riesgo físico, "
             "riesgo de transición, capacidad adaptativa y gobernanza. Es el eje con "
             "comprador naciente y mayor demanda de multilaterales."),
            ("Pensiones (SIPEN)",
             "Dos caras: pulso del sistema (afiliados, rentabilidad, patrimonio) y "
             "scoring por AFP. Solo siete administradoras, así que el panel es pequeño "
             "y la plataforma lo advierte."),
            ("Seguros (SIS)",
             "Solidez financiera de las aseguradoras comerciales y de las ARS de salud, "
             "sobre SIS y SISALRIL. Revisado por una actuaria externa."),
            ("Estructura de la Economía",
             "Vista descriptiva del PIB por sectores de origen con contribución al "
             "crecimiento; se verifica contablemente contra el crecimiento del valor "
             "agregado. Producto para institucionales."),
            ("Desarrollo Social",
             "El único eje sub-nacional: 10 regiones de desarrollo y 32 provincias, con "
             "distribución además de promedio. Sobre SISDOM, MINERD, SIUBEN, ONE y "
             "Banco Mundial."),
            ("Evaluación de Leyes",
             "El eje más original del catálogo: el sujeto es un instrumento normativo. "
             "Evalúa si una ley cumple las metas que ella misma se fijó, contrastando "
             "cada meta contra la fuente oficial del propio Estado evaluado. Dos "
             "expedientes cargados hoy: la Estrategia Nacional de Desarrollo 2030 "
             "(Ley 1-12, 90 indicadores) y Meta RD 2036 (Decreto 337-24). Agregar una "
             "ley es agregar una carpeta de datos, no tocar el motor."),
        ],
        widths=[1.25, 4.55], size=8.0)


def seccion_4(doc: Doc):
    doc.h1("Los tres niveles y la anatomía del informe",
           eyebrow="§4 · Estructura de producto")
    doc.p("Los dieciséis ejes comparten exactamente tres niveles comerciales. La "
          "granularidad —no la extensión— es lo que los separa.")

    doc.table(
        ["Nivel", "Granularidad", "Contenido", "Cadencia", "Audiencia", "Págs."],
        [("Pulse", "Sistema, SIN nombrar",
          "Distribución en bandas, tendencias, comentario del sistema",
          "Periódico / abierto", "Prensa, gremios, tráfico", "1–2 págs"),
         ("Insight", "Entidad o segmento NOMBRADO",
          "Score y perspectiva, radar de pilares, indicadores, pares, narrativa, alertas",
          "Recurrente", "Cliente / comité", "4–6 págs"),
         ("Deep Dive", "A medida",
          "Todo lo de Insight + escenarios + recomendaciones priorizadas + limitaciones",
          "On-demand", "Comité / contraparte", "8–15 págs")],
        widths=[0.72, 1.05, 1.85, 0.82, 0.92, 0.5], size=7.8)

    doc.callout(
        "Regla no negociable del nivel abierto",
        "Pulse NUNCA emite un identificador de entidad. Hay un sensor de anonimización "
        "por sector que lo verifica. Es lo que permite publicar el nivel abierto sin "
        "exponer a ninguna entidad supervisada, y es también lo que hace que el salto a "
        "Insight tenga valor: pagar es pasar de la banda al nombre.",
        tone="accent")

    doc.h2("4.1 Anatomía canónica del informe")
    doc.p("Una sola anatomía, escalada por nivel. Las tres salidas —en pantalla, PDF y "
          "Word— se ensamblan desde una única estructura de contenido, así que no pueden "
          "divergir.")
    doc.table(
        ["#", "Sección", "Qué hace", "Pulse", "Insight", "Deep"],
        [("1", "Portada", "Identidad, sujeto, período, veredicto de titular", "sí", "sí", "sí"),
         ("2", "Índice", "Navegación", "—", "—", "sí"),
         ("3", "Resumen ejecutivo",
          "EL VEREDICTO: tesis, cifra clave, lectura. Decisión extraíble en 30 segundos",
          "sí", "sí", "sí"),
         ("4", "Contexto", "Encuadre del mercado; por qué importa", "—", "sí", "sí"),
         ("5", "Hallazgos y análisis",
          "Dimensiones con peso y aporte, drivers vs. lastres, posición relativa",
          "—", "sí", "sí"),
         ("6", "Recomendaciones",
          "Priorizadas en Inmediato / Corto / Mediano, por audiencia", "—", "—", "sí"),
         ("7", "Metodología y fuentes",
          "Qué mide, con qué pesos, qué fuentes, cadencia, cobertura, versión",
          "mini", "sí", "sí"),
         ("8", "Limitaciones y calidad del dato",
          "Brechas, supuestos de rúbrica, ausencia de backtest, qué NO cubre",
          "—", "sí", "sí"),
         ("9", "Glosario", "Términos e índice", "—", "—", "sí"),
         ("10", "Fuentes y referencias",
          "Citas con fuente · licencia · fecha de consulta · URL", "—", "—", "sí")],
        widths=[0.24, 1.05, 2.35, 0.52, 0.58, 0.46], size=7.8)

    doc.rich([("El detalle que hay que destacar en la presentación. ",
               {"bold": True, "color": INK}),
              ("Las secciones 7, 8 y 10 —Metodología, Limitaciones y Fuentes— son las "
               "más respetadas y peor ejecutadas del mercado, y aquí se generan "
               "automáticamente a partir de la procedencia real de cada variable. No las "
               "redacta nadie. Esa es la demostración más concreta del foso, y merece su "
               "propia lámina.", {})])

    doc.h2("4.2 Elementos visuales del informe")
    doc.bullets([
        ("Cita destacada: ", "texto grande en acento con barra vertical, para la cifra o "
         "el hallazgo clave."),
        ("Llamada de cifra: ", "número grande con etiqueta, estilo infografía."),
        ("Figuras numeradas con leyenda ", "(«Figura N: …»)."),
        ("Tablas: ", "cabecera en tinta, filas alternas, cifras tabulares."),
        ("Encabezado corrido en cada página ", "con producto y sujeto, más paginación."),
        ("Marca de agua por nivel: ", "«Vista abierta», «Suscripción», «On-demand»; y "
         "«MUESTRA — DATA ILUSTRATIVA» en las muestras comerciales."),
        ("Cierre con el descargo de SDQ Consulting.", ""),
    ])


def seccion_5(doc: Doc):
    doc.h1("Capacidades transversales", eyebrow="§5 · El resto del producto")
    doc.p("Los ejes son la materia; esto es la maquinaria. Buena parte del valor —y del "
          "argumento de defensa competitiva— vive acá, y suele quedar fuera de los "
          "materiales de venta por ser invisible. No debería.")

    doc.h2("5.1 Cerebro de Insights — la narrativa con guardas")
    doc.p("La lectura de analista la escribe un modelo de lenguaje, pero circunscrito. "
          "El sistema le pasa la identidad de la casa, la doctrina del eje, el estándar "
          "epistémico y el marco de audiencia; y después verifica lo que escribió.")
    doc.bullets([
        ("Guardrail numérico: ", "recomputa toda cifra antes de publicar y elimina la que "
         "no sea trazable."),
        ("Regla del sujeto: ", "toda clave de cuota o concentración nombra su población. "
         "Nació de un error real —se publicó «cuatro compañías concentran el 87,1%» "
         "cuando eran cuatro ramos— y hoy la vigila un test."),
        ("Relaciones computadas: ", "dirección, deltas, superlativos y posiciones se "
         "calculan en código; el modelo los copia."),
        ("Universo comparable: ", "un score parcial no rankea contra uno completo, ni en "
         "el contexto del modelo ni en la tabla renderizada."),
        ("Marcos de audiencia: ", "el mismo hallazgo se redacta distinto para un "
         "inversionista, un gobierno, una empresa o un multilateral."),
        ("Tres idiomas: ", "español, inglés y francés."),
    ])

    doc.h2("5.2 Alertas accionables — el producto de flujo")
    doc.p("La categoría que convierte la plataforma de consulta en suscripción. No es un "
          "monitor de documentos: es el aviso de que una cifra que el cliente eligió "
          "vigilar cambió de manera que importa.")
    doc.p("Seis disparadores: cruce de umbral declarado, cambio de banda, movimiento en "
          "el ranking del universo comparable, pérdida (o recuperación) del dato que "
          "sostenía una dimensión, orfandad del insumo de una credencial de validación, "
          "y edición nueva de una fuente recurrente.")
    doc.p("Y —lo que la hace SDQ— seis vetos antes de entregar:")
    doc.table(
        ["#", "Veto", "Regla"],
        [("1", "Frescura", "Verificada vigente publica; obsoleta no; e indeterminada, tampoco."),
         ("2", "Brecha", "Ninguna entrada ausente produce alerta; la ausencia usa su propio disparador."),
         ("3", "Comparabilidad", "Rankings y superlativos solo dentro del universo comparable."),
         ("4", "Sujeto", "Toda cuota o concentración nombra su población en el texto."),
         ("5", "Relación", "Dirección, delta y posición vienen del cómputo, no de la prosa."),
         ("6", "Vocabulario", "Nada de «predice», «probabilidad de» ni «riesgo de quiebra».")],
        widths=[0.3, 1.0, 4.5])
    doc.p("Lo vetado se lista, no desaparece: la bandeja del cliente dice «3 señales "
          "retenidas por frescura del dato». Un veto silencioso se leería como que no "
          "pasó nada. Canales: en la aplicación, correo y webhook firmado (HMAC-SHA256, "
          "con autodesactivación tras fallos repetidos).")

    doc.h2("5.3 SDQ Data API — el canal máquina a máquina")
    doc.p("Para que un cliente incorpore los activos de SDQ a sus propios modelos de "
          "riesgo, tableros y comités. No es un espejo del dato público: eso ya lo "
          "publican las fuentes y no es un producto.")
    doc.table(
        ["Endpoint", "Qué sirve"],
        [("GET /catalog", "Manifiesto de exposición, auto-extensible."),
         ("GET /catalog/changes", "Qué cambió desde el último corte."),
         ("GET /series", "Series canónicas normalizadas con período, unidad, frecuencia y linaje."),
         ("GET /scores/{sector}", "Desglose dimensional numérico del score. Jamás narrativa."),
         ("GET /signals/{sector}", "Señales del motor determinista."),
         ("GET /quality/{sector}", "El registro de honestidad servido al cliente."),
         ("GET /forecasts/{sector}", "Pronósticos donde el eje los produce.")],
        widths=[1.5, 4.3], mono_cols=(0,))
    doc.p("Cuatro exclusiones, y solo esas: el payload crudo del conector, la narrativa "
          "completa del informe (servirla por API canibalizaría el producto de reporte), "
          "el núcleo de propiedad intelectual —prompts, doctrina, pesos del modelo "
          "entrenado, rúbricas internas— y todo activo cuya licencia de origen no conste. "
          "Las llaves tienen cuota y bitácora de uso.")

    doc.h2("5.4 Research a Medida — la pregunta libre")
    doc.p("El comprador escribe una pregunta en lenguaje natural en vez de elegir una "
          "plantilla. El motor resuelve qué ejes y qué entidad nombra, trae el resultado "
          "ya computado de esos motores, descompone la pregunta en sub-preguntas con "
          "metodología y procedencia, y escribe la respuesta usando SOLO ese dato. "
          "Lo que ningún motor computa —una proyección a futuro— se declara brecha, no se "
          "rellena. Es el ancla de precio más alta del catálogo.")

    doc.h2("5.5 Deal Scoring")
    doc.p("Índice de atractivo de una operación concreta, sobre una rúbrica de siete ejes "
          "anclada a los índices reales de la plataforma: el atractivo del sector alimenta "
          "la validación de mercado, el riesgo país alimenta la preparación regulatoria y "
          "la resiliencia climática alimenta el factor de clima. Se auto-rotula como "
          "rúbrica declarada, no como modelo entrenado.")

    doc.h2("5.6 Contexto de Marca (Brand Intel)")
    doc.p("El único módulo con datos privados de cliente: cada encargo está aislado por "
          "identificador y esa frontera se impone en la capa de datos, no se confía al "
          "que llama. Motores de atribución, categoría, deflactación, embudo, segmentos, "
          "significancia estadística, pronóstico, escenarios y vigilancia. Donde falta un "
          "insumo, el resultado es ausencia con motivo declarado — que es lo que permite "
          "distinguir «no hubo movimiento» de «no se midió».")

    doc.h2("5.7 Herramientas y superficies de plataforma")
    doc.table(
        ["Superficie", "Qué es"],
        [("Resumen ejecutivo", "Vista transversal de todos los ejes activos."),
         ("Comparador", "Comparación de sujetos entre ejes con lectura narrativa comparativa."),
         ("Market Brief", "Sesión informativa periódica generada sobre el corte vigente."),
         ("Metodología", "Cómo se mide cada índice, con pesos y fuentes, servido al cliente."),
         ("Catálogo de productos", "La tienda: qué hay, qué nivel, qué incluye."),
         ("Mi plan", "Qué compró el cliente y hasta cuándo."),
         ("Mis vigilancias", "La watchlist de alertas."),
         ("Búsqueda global", "Paleta de comandos con ⌘K sobre entidades, sectores y países."),
         ("Monitor de productos", "Interno: readiness y activación pública por celda."),
         ("Inteligencia de Fuentes", "Interno: sugerir, evaluar e integrar fuentes nuevas."),
         ("Consola de operaciones", "Interno: 91 operaciones con cadencia, historial y reintento.")],
        widths=[1.3, 4.5])

    doc.h2("5.8 El gate de publicación (readiness G1–G5)")
    doc.p("Un producto puede estar cableado y funcionando internamente sin estar "
          "disponible para el público. La activación es siempre manual y explícita, y el "
          "monitor la bloquea si el producto no cruza el umbral de calidad de su nivel.")
    doc.table(
        ["Gate", "Qué mide", "Peso"],
        [("G1 · Datos", "Ingesta de la fuente autoritativa operativa y fresca", "30%"),
         ("G2 · Motor", "Índice explicable con el scoring corriendo", "25%"),
         ("G3 · Narrativa", "Marco SCQA operativo y guarda anti-alucinación en verde", "15%"),
         ("G4 · Plantilla", "El reporte del nivel renderiza", "15%"),
         ("G5 · Validación", "Estado de validación y doctrina firmados", "15%")],
        widths=[1.1, 4.0, 0.7])
    doc.p("Umbral de activación: 0,75 para Pulse y 0,85 para Insight y Deep Dive — más "
          "exigente donde se nombra una entidad, porque ahí está la reputación.")

    doc.h2("5.9 Acceso, roles y facturación")
    doc.bullets([
        ("Dos ejes de permiso independientes. ",
         "El ROL (viewer · analyst · admin · super_admin) gobierna qué puede hacer un "
         "usuario; el NIVEL DE ACCESO (free · pro · enterprise) gobierna qué contenido "
         "ve. No se mezclan."),
        ("Acceso compuesto: ", "nivel manual, suscripción vigente o derecho otorgado — "
         "cualquiera de los tres concede, sin mutar los otros."),
        ("Pasarela: ", "PayPal, con planes por SKU e intervalo, sincronización de planes "
         "y webhook idempotente del ciclo de vida."),
        ("Fiscalidad dominicana resuelta: ", "desglose de ITBIS, secuencias de "
         "comprobantes, comprobante fiscal electrónico (e-CF) de la DGII con su XML, "
         "notas de crédito y reporte 607. Es una barrera de entrada real para un "
         "competidor extranjero."),
    ])


def seccion_6(doc: Doc, skus):
    doc.h1("Modelo comercial", eyebrow="§6 · Cómo se cobra")

    doc.h2("6.1 La estructura de SKU")
    doc.p("Un SKU es lo que se cobra, independiente del proveedor de pago. Cinco familias "
          "cubren todo el catálogo, y cada una concede acceso a uno o más pares "
          "(sector, nivel) que el control de acceso compone. Un derecho de nivel N cubre "
          "los niveles inferiores del mismo alcance.")
    doc.table(
        ["Familia de SKU", "Qué concede", "Intervalo"],
        [("insight:{sector}", "El Insight de UN eje", "mensual · anual"),
         ("deep_dive:{sector}", "Un Deep Dive de UN eje", "una vez"),
         ("all_access", "El Insight de TODOS los ejes", "mensual · anual"),
         ("enterprise", "Catálogo completo: Insight y Deep Dive de todos los ejes",
          "mensual · anual"),
         ("special:{slug}", "Informe especial cotizado a medida", "una vez")],
        widths=[1.4, 3.4, 1.0], mono_cols=(0,))
    doc.rich([
        ("Tamaño del tarifario hoy: ", {"bold": True, "color": INK}),
        (f"{len(skus)} SKU canónicos vendibles — {len(skus) - 2} por producto "
         f"(cada eje con su Insight y su Deep Dive) más los dos bundles. Los informes "
         f"especiales no se enumeran porque son a medida.", {}),
    ])

    doc.h2("6.2 Tarifario — PROPUESTA")
    doc.callout(
        "Estado del precio",
        "Todos los importes de abajo son PROPUESTA para calibración del dueño, salvo el "
        "ancla del Research a Medida (US$3,500 por encargo, también provisional). Nada "
        "rige hasta publicarse en la consola de tarifario. En la presentación comercial: "
        "no imprima precios cerrados sin confirmarlos con Ricardo Mercado.",
        tone="warn")
    doc.table(
        ["Producto", "SKU", "Intervalo", "Precio propuesto", "Nota"],
        [(a, b, c, d, e) for a, b, c, d, e in PRECIOS],
        widths=[1.15, 1.25, 0.75, 0.95, 1.7], mono_cols=(1, 3), size=7.8)
    doc.p("Referencia de mercado para el encuadre: un encargo de consultoría tradicional "
          "equivalente cuesta entre US$7.000 y US$25.000 o más, y no entrega ni "
          "procedencia por variable ni actualización.")

    doc.h2("6.3 Estado operativo del cobro")
    doc.p("Para que la presentación no prometa un flujo que hoy no está encendido:")
    doc.table(
        ["Pieza", "Estado"],
        [("Código de checkout y suscripción", "Listo; ambiente sandbox y producción por configuración."),
         ("Desglose de ITBIS y facturación", "Listo."),
         ("Comprobante fiscal electrónico (e-CF)", "Implementado."),
         ("Credenciales de la pasarela", "Pendientes del dueño: falta el secreto y activar la bandera."),
         ("Planes de cobro por SKU e intervalo", "Pendientes de crear en el panel del proveedor."),
         ("Tarifario publicado", "Pendiente: sin precios publicados no hay venta posible.")],
        widths=[2.0, 3.8])

    doc.h2("6.4 El argumento de la escalera")
    doc.p("La estructura de tres niveles no es solo empaquetado; es el recorrido comercial "
          "y conviene dibujarlo así en la presentación.")
    doc.bullets([
        ("Pulse abre el mercado. ", "Es publicable, citable por la prensa y no expone a "
         "ninguna entidad. Genera la reputación de la casa y el tráfico."),
        ("Insight monetiza el nombre. ", "El salto de la banda a la entidad nombrada es "
         "exactamente lo que el cliente institucional necesita y no puede publicar solo."),
        ("Deep Dive y Research capturan el encargo. ",
         "Es donde se compite contra la consultoría tradicional, con una fracción del "
         "precio y con procedencia auditable que la consultoría no entrega."),
        ("Las alertas y la API retienen. ",
         "Convierten una compra puntual en una dependencia operativa: cuando el modelo de "
         "riesgo del cliente consume la serie de SDQ, el costo de cambio deja de ser el "
         "precio."),
    ])


def seccion_7(doc: Doc):
    doc.h1("Qué se puede afirmar y qué no", eyebrow="§7 · Filtro obligatorio de mensaje")
    doc.callout(
        "Por qué esta sección existe",
        "Un comprador institucional puede auditar. El mayor riesgo reputacional de "
        "SDQ·MIP es una brecha entre el material de venta y lo que el código sostiene. "
        "Toda pieza de la presentación pasa por acá antes de salir.",
        tone="alert")

    doc.h2("7.1 La regla madre")
    doc.p("En cualquier material se distinguen siempre tres cosas, y no se confunden:")
    doc.table(
        ["Categoría", "Qué es", "Cómo se rotula"],
        [("Dato medido", "Viene de una fuente externa verificable", "live"),
         ("Juicio experto", "Valor de criterio de la casa, declarado, fechado y atribuido",
          "rubric"),
         ("Modelo", "O es ENTRENADO —aprende de un desenlace externo y tiene backtest "
                    "honesto— o es un ÍNDICE EXPLICABLE (fórmula transparente). De un "
                    "índice no se dice «predictivo».", "según el caso")],
        widths=[1.0, 3.8, 1.0], mono_cols=(2,))

    doc.h2("7.2 Reglas duras por producto")
    doc.table(
        ["Producto", "SE PUEDE decir", "NO se puede decir"],
        [("Rating bancario",
          "«Scoring de entidad explicable y auditable, con evidencia por eje y procedencia "
          "por variable». «Aproximador del método SDQ que preserva la explicabilidad».",
          "«Modelo predictivo de default o quiebra». El modelo aprende a reproducir la "
          "rúbrica determinista; su precisión mide fidelidad a la fórmula, no poder "
          "predictivo. El foso es la explicabilidad, no el aprendizaje automático."),
         ("Política monetaria / TPM",
          "«Modelo predictivo con backtest honesto», «expanding-window one-step-ahead», "
          "«panel point-in-time», «track record en vivo».",
          "«Track record probado». El registro en vivo arranca vacío y crece a razón de "
          "un dato por reunión: se dice «en construcción». Declarar el leakage residual "
          "conocido."),
         ("Deal Scoring",
          "«Índice de atractivo explicable», «rúbrica de siete ejes anclada a índices "
          "reales», «propiedad intelectual metodológica de la casa».",
          "«Probabilidad de éxito» ni «modelo predictivo». Es rúbrica declarada y el "
          "código se auto-rotula como no entrenado."),
         ("IRMP (riesgo macro-político)",
          "«Índice de riesgo país sobre dato real de gobernanza, macro, calificación "
          "soberana y eventos, MÁS juicio experto declarado en tres variables "
          "institucionales».",
          "«100% dato» ni «cero rúbrica». Tres insumos son juicio experto y pesan cerca "
          "de un tercio. Es legítimo y estándar en rating; presentarlo como medición pura "
          "no lo es."),
         ("Índices sectoriales, seguros, pensiones, ESG, social",
          "«Compuesto de dato real con procedencia por variable». «Las dimensiones sin "
          "dato se declaran brecha y se excluyen, no se rellenan».",
          "Ocultar los rótulos de rúbrica donde los hay. El material debe ser consistente "
          "con lo que la interfaz ya muestra.")],
        widths=[1.05, 2.4, 2.35], size=7.6)

    doc.h2("7.3 Los seis grupos de validación (y por qué no se mezclan)")
    doc.p("«Tiene validación» abarca cosas que no sostienen el mismo argumento de venta. "
          "La plataforma las separa en seis grupos y el material comercial tiene que "
          "respetar esa separación.")
    doc.table(
        ["Grupo", "Qué autoriza a decir"],
        [("A · evento real",
          "Discriminación contra desenlaces reales de entidades."),
         ("B · concluyente",
          "Discriminación contra un desenlace realizado, con intervalo que no cruza cero."),
         ("C · convergente",
          "Coincide con una medida independiente del mismo período. NO es backtest "
          "temporal y no se vende como tal."),
         ("D · parcial",
          "Metodología exigente con resultado acotado, declarado."),
         ("E · corrida sin credencial a favor",
          "El motor se aplicó y no dejó una afirmación vendible: o el intervalo cruza "
          "cero, o no lo cruza pero está del lado equivocado (la señal ordena invertido, "
          "que es un hallazgo, no una ausencia). Honesto; no es credencial."),
         ("F · sin backtest transversal",
          "Descriptivo, con el obstáculo declarado.")],
        widths=[1.4, 4.4])

    doc.h3("Qué clase de validación admite cada eje")
    doc.p("Esto es un hecho de DISEÑO —si el eje tiene motor de validación, contra qué "
          "desenlace, y qué lo impide cuando no lo tiene— y por eso sí se puede escribir. "
          "El VEREDICTO de la última corrida no: ese envejece y se lee en vivo.",
          space_after=4)
    doc.table(
        ["Eje", "Estado estructural de validación"],
        [(name, EJES[key][2]) for key, name in _EJE_NOMBRES],
        widths=[1.4, 4.4], size=7.8)

    doc.h2("7.4 La regla de oro: ninguna cifra de validación se escribe a mano")
    doc.callout(
        "Esto no es una preferencia de estilo",
        "Producción sirvió durante diecinueve días una cifra de discriminación calculada "
        "con un score que ya no existía, mientras el material de venta citaba otra. La "
        "cura fue estructural: hay UNA sola fuente de la tabla comercial de validación, "
        "que computa la afirmación leyendo el reporte vigente de cada motor y veta lo que "
        "no puede verificar como fresco.",
        tone="alert")
    doc.p("En términos prácticos para quien arma la presentación:")
    doc.bullets([
        ("No copie ninguna cifra de validación de este documento. ",
         "No hay ninguna, a propósito."),
        ("Pídala en el momento de maquetar. ",
         "Se lee de la propia plataforma, en el punto de credenciales de producto, y "
         "viene con su población al lado y su veredicto de frescura."),
        ("Si una cifra viene vetada, no la use ni la esconda. ",
         "Las vetadas se listan con su motivo: un veto silencioso se lee como que el eje "
         "no tiene validación."),
        ("Y si quiere saber cómo está un eje hoy, pregúntele a la plataforma. ",
         "Cualquier tabla escrita ya está vieja — incluida cualquier tabla de este "
         "documento que alguien copie a mano en el futuro."),
    ])
    doc.p()
    doc.h3("Tres correcciones que el material anterior tenía mal")
    doc.bullets([
        ("El denominador de la cohorte bancaria. ",
         "Del universo de quiebras históricas del ledger, solo una parte es evaluable: el "
         "resto no puede disparar por construcción, porque el indicador que la regla exige "
         "no existía en ese período. Citar el total infla el denominador sin agregar "
         "evidencia."),
        ("El desenlace del backtest bancario no es «quiebras». ",
         "Es distress financiero, y está compuesto en proporciones muy desiguales de "
         "pérdidas sostenidas, deterioro de crédito y solvencia — esta última nunca "
         "disparó. Citar el agregado sin decir de qué está hecho es la afirmación más "
         "frágil del catálogo."),
        ("Sobre QUIÉN se midió. ",
         "El panel bancario no es «bancos»: es todo el universo supervisado por la SIB, y "
         "casi la mitad de las observaciones son entidades de intermediación cambiaria y "
         "fiduciarias, que no otorgan crédito. Están por diseño del producto, que es un "
         "score de entidad financiera y no solo de banca — pero una cifra de "
         "discriminación sola se lee como si fuera entre bancos, y no es lo que se midió."),
    ])


def seccion_8(doc: Doc):
    doc.h1("Brief para el diseñador", eyebrow="§8 · Encargo")

    doc.h2("8.1 Qué se espera")
    doc.p("Una presentación comercial de SDQ·MIP que un vendedor pueda usar frente a un "
          "comité institucional, y de la que se puedan extraer piezas sueltas (una lámina "
          "de un eje, la lámina de precios) sin que pierdan sentido.")
    doc.bullets([
        ("Formato: ", "16:9. Versión maestra en claro; portada y cierre en oscuro."),
        ("Extensión sugerida: ", "22 a 28 láminas para la versión completa, más un corte "
         "corto de 10 para primera reunión."),
        ("Entregables: ", "archivo editable con maestras y componentes, PDF de alta, y "
         "un juego de láminas sueltas por eje para armar propuestas a medida."),
        ("Sistema, no colección: ", "cada tipo de lámina debería existir como maestra "
         "reutilizable, porque el catálogo crece."),
    ])

    doc.h2("8.2 Guion sugerido, lámina por lámina")
    doc.table(
        ["#", "Lámina", "Contenido", "Fuente en este documento"],
        [("1", "Portada", "Bloque de marca sobre oscuro; una línea de posicionamiento", "§1"),
         ("2", "El problema", "El dato existe y está fragmentado; el costo de reconstruirlo", "§1.1"),
         ("3", "La respuesta", "La frase de una línea, en grande", "§1"),
         ("4", "El foso", "Las cuatro reglas de procedencia, como sistema visual", "§1.2"),
         ("5", "La arquitectura", "La figura 1, animada por capas", "§1.4"),
         ("6", "El catálogo", "Los 16 ejes como retícula de fichas", "§3"),
         ("7–10", "Ejes destacados", "Cuatro fichas ampliadas: banca, macro, social, leyes", "§3.1"),
         ("11", "Los tres niveles", "La escalera Pulse → Insight → Deep Dive", "§4"),
         ("12", "Anatomía del informe", "La tabla de secciones por nivel, visual", "§4.1"),
         ("13", "Metodología automática", "La lámina del foso concreto", "§4.1, nota"),
         ("14", "Alertas", "Los seis disparadores y los seis vetos", "§5.2"),
         ("15", "Data API", "El canal máquina a máquina y las cuatro exclusiones", "§5.3"),
         ("16", "Research a medida", "La pregunta libre con gate de honestidad", "§5.4"),
         ("17", "Herramientas", "Deal Scoring, Contexto de Marca, Comparador", "§5.5–5.7"),
         ("18", "El gate de publicación", "G1–G5 y por qué existe", "§5.8"),
         ("19", "Honestidad como producto", "Dato / rúbrica / brecha, con ejemplo real", "§7.1"),
         ("20", "Audiencias", "Quién compra qué", "§1.3"),
         ("21", "Planes y precios", "Los tres niveles con el tarifario", "§6.2"),
         ("22", "Por qué ahora", "Fiscalidad local resuelta, catálogo completo, API viva", "§5.9"),
         ("23", "Cierre", "Bloque de marca, contacto, descargo", "—")],
        widths=[0.35, 1.15, 2.75, 1.05], size=7.6)

    doc.h2("8.3 Aciertos y errores previsibles")
    doc.table(
        ["Haga esto", "No haga esto"],
        [("Use el medidor de arco como motivo gráfico recurrente: es el logotipo y es el "
          "producto.",
          "No invente iconografía sectorial genérica (una fábrica, una palmera). El "
          "sistema es tipográfico y de dato."),
         ("Muestre pantallas reales del producto y páginas reales del informe.",
          "No maquete pantallas ficticias más bonitas que las reales: el comprador las "
          "verá en la demo."),
         ("Deje respirar. Cifras grandes, mucho blanco, una idea por lámina.",
          "No llene una lámina con seis estadísticas de relleno. Cada elemento gana su lugar."),
         ("Etiquete la procedencia cuando muestre una cifra, aunque sea de ejemplo.",
          "No use degradados decorativos, sombras dramáticas ni emoji."),
         ("Marque «datos ilustrativos» donde la cifra sea de ejemplo.",
          "No presente una cifra de ejemplo como si fuera un resultado medido."),
         ("Use la barra vertical de acento para citas destacadas.",
          "No use tarjetas con borde de acento a la izquierda como decoración general.")],
        widths=[2.9, 2.9], size=8.0)

    doc.h2("8.4 Assets adjuntos")
    doc.table(
        ["Archivo", "Qué es"],
        [("logo_arco_1024.png", "El símbolo canónico, 1024 px, fondo transparente."),
         ("arco.svg", "El mismo, vectorial — generado desde el código, no dibujado."),
         ("logo_variants.png", "Las dos variantes retiradas frente a la canónica (figura 2)."),
         ("lockup_light.png / lockup_dark.png", "Bloque de marca en claro y en oscuro."),
         ("arch.png", "Diagrama de arquitectura (figura 1)."),
         ("ui_light.png / ui_dark.png", "Muestra de componentes en ambos temas (figura 4)."),
         ("tokens.css", "Los tokens completos, listos para pegar en una herramienta de diseño.")],
        widths=[1.9, 3.9], mono_cols=(0,))

    doc.h2("8.5 Decisiones que el diseñador no puede tomar solo")
    doc.p("Para Ricardo Mercado, antes de que la presentación se cierre. La geometría "
      "del símbolo y la unificación de la paleta ya se resolvieron y salieron del "
      "listado.")
    doc.table(
        ["#", "Decisión", "Recomendación de este documento"],
        [("1", "Confirmar si Arco pasa de identidad de producto a identidad de casa.",
          "Fuera del alcance de este encargo; decidir aparte."),
         ("2", "Cerrar el tarifario.",
          "Sin precios publicados no hay venta; los de §6.2 son propuesta."),
         ("3", "Encender la pasarela de pago.",
          "Tres pasos pendientes del dueño, listados en §6.3."),
         ("4", "Qué ejes se muestran como activos en la presentación.",
          "Solo los que el monitor tenga activados al momento de imprimir.")],
        widths=[0.3, 2.3, 3.2], size=8.0)


def anexos(doc: Doc, catalogo):
    doc.h1("Anexos", eyebrow="§9 · Referencia")

    doc.h2("A · Tokens completos")
    doc.p("Pegue esto en la herramienta de diseño. Es literalmente el archivo que consume "
          "la aplicación, y su espejo en el backend —`shared/brand/tokens.py`— es el que "
          "pintan los informes. Un test compara los dos, así que no pueden divergir.")
    doc.code([
        ":root {",
        f"  --canvas:#{CANVAS};  --surface:#{SURFACE};  --surface-2:#{SURFACE2};",
        f"  --ink:#{INK};     --body:#{BODY};     --muted:#{MUTED};   --faint:#{FAINT};",
        f"  --border:#{BORDER};  --border-strong:#{BORDER_STRONG};",
        f"  --accent:#{ACCENT};  --accent-hover:#{ACCENT_HOVER};  "
        f"--accent-soft:#{ACCENT_SOFT};  --accent-ink:#{ACCENT_INK};",
        f"  --teal:#{TEAL};    --teal-soft:#{TEAL_SOFT};",
        f"  --ok:#{OK};   --ok-soft:#{OK_SOFT};",
        f"  --warn:#{WARN}; --warn-soft:#{WARN_SOFT};",
        f"  --alert:#{ALERT}; --alert-soft:#{ALERT_SOFT};",
        f"  --c1:#{C1}; --c2:#{C2}; --c3:#{C3}; --c4:#{C4}; --c5:#{C5}; --c6:#{C6};",
        f"  --grid:#{GRID};",
        "  --shadow-card: 0 1px 2px rgba(10,26,58,.05), 0 10px 24px -18px rgba(10,26,58,.22);",
        "  --shadow-pop:  0 8px 28px -8px rgba(10,26,58,.22), 0 2px 8px -2px rgba(10,26,58,.12);",
        "}",
        ".dark {",
        f"  --canvas:#{DARK['--canvas']};  --surface:#{DARK['--surface']};  "
        f"--surface-2:#{DARK['--surface-2']};",
        f"  --ink:#{DARK['--ink']};     --body:#{DARK['--body']};     "
        f"--muted:#{DARK['--muted']};   --faint:#{DARK['--faint']};",
        f"  --border:#{DARK['--border']};  --border-strong:#{DARK['--border-strong']};",
        f"  --accent:#{DARK['--accent']};  --accent-hover:#{DARK['--accent-hover']};  "
        f"--accent-soft:rgba(59,130,246,.15);  --accent-ink:#{DARK['--accent-ink']};",
        f"  --teal:#{DARK['--teal']};    --teal-soft:rgba(45,212,191,.14);",
        f"  --ok:#{DARK['--ok']};   --ok-soft:rgba(52,211,153,.14);",
        f"  --warn:#{DARK['--warn']}; --warn-soft:rgba(251,191,36,.14);",
        f"  --alert:#{DARK['--alert']}; --alert-soft:rgba(242,100,90,.15);",
        f"  --c1:#{DARK['--c1']}; --c2:#{DARK['--c2']}; --c3:#{DARK['--c3']}; "
        f"--c4:#{DARK['--c4']}; --c5:#{DARK['--c5']}; --c6:#{DARK['--c6']};",
        f"  --grid:#{DARK['--grid']};",
        "}",
    ])

    doc.h2("B · El símbolo en vectorial")
    doc.p("El arco abre en la parte superior —un hueco de 90°— y el punto flota separado "
          "sobre su centro: es lo que hace legible la metáfora del medidor. Se declara "
          "como trazado con extremos explícitos y no como círculo punteado, porque un "
          "desfase de guion mueve el punto de contacto sin que nadie lo note. Este bloque "
          "se genera desde el código, no se dibuja a mano.")
    from shared.brand import arco_svg
    svg = arco_svg(size=32)
    doc.code([svg[:svg.index("><rect") + 1]]
             + [f"  {t}>" for t in svg[svg.index("<rect"):-len("</svg>")].split(">") if t]
             + ["</svg>"])

    doc.h2("C · Mapa de navegación de la plataforma")
    doc.p("Cinco grupos en la barra lateral. Útil si la presentación incluye una lámina "
          "de recorrido de producto.")
    doc.table(
        ["Grupo", "Elementos"],
        [("Ejes de inteligencia",
          "Financiero · Macroeconómico · Sectorial · Regulatorio y político · Social y "
          "desarrollo · Comercio exterior · ESG y clima · Pensiones · Seguros"),
         ("Datos",
          "Banca (SIB) · Operaciones · Macro (BCRD) · Social (ONE) · Comercio (DGA) · "
          "Gobernanza (WGI) · Pensiones (SIPEN)"),
         ("Herramientas",
          "Research a Medida · Contexto de Marca · Deal Scoring · Market Brief"),
         ("Plataforma",
          "Resumen ejecutivo · Comparador · Metodología · Catálogo de productos · Mi plan · "
          "Mis vigilancias · Monitor de productos · Inteligencia de Fuentes · Configuración"),
         ("Administración",
          "Usuarios · Tarifario · Pagos · Comprobantes fiscales")],
        widths=[1.2, 4.6])

    doc.h2("D · Glosario de siglas")
    doc.table(
        ["Sigla", "Qué es"],
        [("Perfil SDQ", "Lectura financiera de dos ejes: Ejecución y Resiliencia."),
         ("ISF", "Índice de Solidez Financiera (seguros)."),
         ("ISA", "Índice de Solidez de la AFP (pensiones)."),
         ("IRMP", "Índice de Riesgo Macro-Político (país)."),
         ("IAI", "Índice de Atractivo de Inversión (sector)."),
         ("SGPS", "Grado-potencial sectorial, por letra A–D."),
         ("IRC", "Índice de Resiliencia Climática (país)."),
         ("IDM", "Índice de Desarrollo Multidimensional (región / provincia)."),
         ("ITT", "Índice de Tracción Turística."),
         ("IZF", "Índice de Zonas Francas."),
         ("IRSE", "Índice de Resiliencia del Sistema Eléctrico."),
         ("IDT", "Índice de Desarrollo de Telecomunicaciones."),
         ("ICC", "Índice de Coyuntura de la Construcción."),
         ("TPM", "Tasa de Política Monetaria del BCRD."),
         ("SCQA", "Situación · Complicación · Pregunta · Respuesta (marco narrativo)."),
         ("e-CF", "Comprobante Fiscal Electrónico de la DGII."),
         ("Pulse / Insight / Deep Dive", "Los tres niveles comerciales.")],
        widths=[1.1, 4.7])

    doc.h2("E · Fuentes autoritativas por institución")
    doc.table(
        ["Institución", "Qué aporta"],
        [("SIB · SIMBAD", "Estados financieros de todo el universo supervisado."),
         ("BCRD", "Macro, cuentas nacionales, sectores de origen, turismo, inversión "
                  "extranjera, comunicados de política monetaria."),
         ("SIPEN", "Sistema de pensiones y AFP."),
         ("SIS · SISALRIL", "Seguros comerciales y ARS de salud."),
         ("DGA", "Comercio exterior y socios comerciales."),
         ("ONE · MEPyD (SISDOM) · SIUBEN · MINERD", "Estadística social y educativa."),
         ("SIE · CNE · Organismo Coordinador", "Sector eléctrico."),
         ("CNZFE", "Zonas francas."),
         ("INDOTEL · ITU", "Telecomunicaciones."),
         ("MIVHED", "Permisos de construcción."),
         ("DGII", "Contribuyentes y régimen de comprobantes fiscales."),
         ("Banco Mundial (WGI · WDI)", "Gobernanza y desarrollo comparados."),
         ("GDELT", "Tensión de eventos político-institucionales."),
         ("ND-GAIN · HURDAT2 · OWID/EM-DAT", "Exposición y desenlaces climáticos."),
         ("Comtrade · CEPALSTAT · Ember · IPU", "Comparables internacionales.")],
        widths=[1.7, 4.1])

    doc.p()
    _rule(doc.d, color=BORDER_STRONG)
    doc.p("Documento generado desde el propio repositorio de la plataforma "
          "(scripts/build_dossier_comercial.py). El catálogo y el tarifario se leen del "
          "código, así que este archivo se regenera y no se edita a mano. Las cifras de "
          "validación se consultan en vivo — ver §7.4.",
          size=8, color=MUTED, italic=True)


# ─────────────────────────── ensamblado ───────────────────────────

def build(out_dir: str, assets_dir: str) -> str:
    catalogo, skus = load_catalog()
    hoy = date.today().isoformat()

    doc = Doc(assets_dir)
    doc.running_furniture("SDQ·MIP  ·  DOSSIER DE PRODUCTO Y MARCA  ·  HANDOFF DE DISEÑO")

    cover(doc, hoy)
    indice(doc)
    seccion_0(doc)
    seccion_1(doc)
    seccion_2(doc)
    seccion_3(doc, catalogo)
    seccion_4(doc)
    seccion_5(doc)
    seccion_6(doc, skus)
    seccion_7(doc)
    seccion_8(doc)
    anexos(doc, catalogo)

    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, "SDQMIP_Dossier_Comercial_y_Marca.docx")
    doc.d.save(path)
    return path


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out", default="docs/comercial")
    ap.add_argument("--assets", default=None)
    args = ap.parse_args()
    assets = args.assets or os.path.join(args.out, "assets")
    path = build(args.out, assets)
    print(f"OK → {path}")


if __name__ == "__main__":
    main()
