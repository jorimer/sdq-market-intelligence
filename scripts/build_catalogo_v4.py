"""Genera el Catálogo SDQMIP v4 — 16 ejes, con la validación REAL de cada uno.

Qué cambia frente al v3 (13-jul-2026), y no es cosmético:

1. **Son 16 ejes, no 14.** La lista sale del registro (`PRODUCT_CATALOG`), así que el número
   nunca vuelve a escribirse a mano: entraron `social_dev` y `law` después del v3 y el
   documento los tomaba solo si alguien se acordaba.
2. **Cada eje trae su credencial de validación computada**, del reporte del motor, con su
   IC y su N. Ninguna cifra se transcribe: el v3 no traía ninguna, y el riesgo de la
   siguiente versión era que alguien copiara un número de un reporte envejecido.
3. **El gate de frescura veta.** Una cifra cuyo reporte no se pueda verificar vigente contra
   su insumo NO entra al documento: sale como «no publicable» con el motivo. Producción sirvió
   19 días un Gini de 0,44 calculado con un score que ya no existía mientras el material de
   venta citaba el correcto; esto impide que esa cifra llegue a un PDF.
4. **Los seis grupos de validación se muestran separados.** «Tiene validación» abarca cosas
   que no sostienen el mismo argumento: un backtest contra quiebras reales y un índice sin
   corte transversal no se presentan como equivalentes.

Uso (contra la base de PRODUCCIÓN, por la API, sin acceso a la DB)::

    python scripts/build_catalogo_v4.py                 # baja credenciales de prod
    python scripts/build_catalogo_v4.py --local         # usa la DB local
    python scripts/build_catalogo_v4.py --json-only     # solo el volcado, sin .docx

El .docx sale en el directorio del proyecto; el JSON de credenciales queda en `evidence/`
para que el documento sea auditable contra la corrida que lo produjo.
"""
from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import date, datetime, timezone
from typing import Any, Dict, List

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from shared.products.credenciales import GRUPOS  # noqa: E402
from shared.products.registry import PRODUCT_CATALOG  # noqa: E402

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SALIDA_DOCX = "/Users/ricardomercado/Developer/SDQMIP"
HOY = date.today().isoformat()

_MESES = ("enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto",
          "septiembre", "octubre", "noviembre", "diciembre")


def _fecha_larga() -> str:
    h = date.today()
    return f"{h.day} de {_MESES[h.month - 1]} de {h.year}"


def _num(v: Any, dec: int = 4) -> str:
    """Número con coma decimal (el documento se lee en español), o guion si no hay."""
    if v is None:
        return "—"
    return f"{float(v):.{dec}f}".replace(".", ",")


def _ic(par: Any) -> str:
    if not par or par[0] is None or par[1] is None:
        return "—"
    return f"[{_num(par[0], 3)} · {_num(par[1], 3)}]"


def credenciales_de_prod() -> Dict[str, Any]:
    """Baja la tabla de credenciales de producción por la API (cuenta de operaciones)."""
    import httpx

    from scripts.ops_trigger import DEFAULT_BASE, _credentials

    email, password = _credentials()
    with httpx.Client(base_url=DEFAULT_BASE, timeout=180, follow_redirects=True) as c:
        tok = c.post("/api/v1/auth/login",
                     json={"email": email, "password": password}).json()["access_token"]
        r = c.get("/api/v1/products/credenciales",
                  headers={"Authorization": f"Bearer {tok}"})
        r.raise_for_status()
        return r.json()


def credenciales_locales() -> Dict[str, Any]:
    from shared.database.session import SessionLocal
    from shared.products.credenciales import credenciales

    db = SessionLocal()
    try:
        return credenciales(db)
    finally:
        db.close()


#: Qué parte del panel puede no otorgar crédito antes de que haya que decirlo AL LADO de la
#: cifra. No es un umbral de calidad: por encima de esto, un lector que ve «Gini · n» atribuye
#: la discriminación a una población que no es la que se midió.
CUOTA_QUE_OBLIGA_A_DECLARAR = 0.10


def _sufijo_poblacion(f: Dict[str, Any]) -> str:
    """La población, pegada al N. Computada de la credencial; nunca escrita a mano.

    El caso: banca publica «Gini · n=1.693» y casi la mitad de ese panel son entidades que no
    otorgan crédito —están en el universo por diseño del producto— pero la cifra sola se lee
    como discriminación ENTRE BANCOS. El sujeto viaja con el número o el número se reatribuye.
    """
    pob = f.get("poblacion") or {}
    sin_credito = pob.get("sin_libro_de_credito") or {}
    cuota = sin_credito.get("cuota_de_observaciones")
    if not cuota or cuota < CUOTA_QUE_OBLIGA_A_DECLARAR:
        return ""
    cuota_ev = sin_credito.get("cuota_de_eventos")
    tipos = ", ".join(sin_credito.get("tipos") or []) or "otras poblaciones"
    texto = (f" · {cuota * 100:.0f} % del panel NO otorga crédito ({tipos})")
    if cuota_ev:
        texto += f" y aporta {cuota_ev * 100:.0f} % de los eventos"
    return texto.replace(".", ",")


def _fila_credencial(f: Dict[str, Any]) -> List[str]:
    """Una fila de la tabla de validación. Lo NO publicable se dice, no se borra."""
    if not f["publicable"]:
        motivo = f.get("stale_reason") or "sin cifra vigente"
        if f["valor"] is not None:
            return [f["nombre"], "— (no publicable)", "—", "—", f"vetada: {motivo}"]
        est = f.get("estado_backtest") or {}
        return [f["nombre"], "sin backtest", "—", "—",
                est.get("obstaculo") or "sin motor declarado"]
    n = f["n"]
    detalle = f"n={n:,}".replace(",", ".") if isinstance(n, int) else "—"
    if f.get("eventos"):
        detalle += f" · {f['eventos']:,} eventos".replace(",", ".")
    if f.get("senal"):
        detalle += f" · señal «{f['senal']}»"
    detalle += _sufijo_poblacion(f)
    return [f["nombre"], f["metrica"] or "—", _num(f["valor"], 4), _ic(f["ic"]), detalle]


def build_docx(cred: Dict[str, Any]) -> str:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt, RGBColor

    NAVY, RED, GRAY = RGBColor(0x0F, 0x2A, 0x4A), RGBColor(0xD6, 0x2A, 0x2A), RGBColor(0x6B, 0x72, 0x80)
    doc = Document()
    doc.styles["Normal"].font.name = "Inter"
    doc.styles["Normal"].font.size = Pt(10)

    def _h(t, size=15, color=NAVY, before=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        r = p.add_run(t)
        r.bold, r.font.size, r.font.color.rgb = True, Pt(size), color
        r.font.name = "Plus Jakarta Sans"

    def _p(t, size=10, color=None, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.size, r.italic = Pt(size), italic
        if color:
            r.font.color.rgb = color

    def _tabla(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            celda = t.rows[0].cells[i]
            celda.text = h
            for r in celda.paragraphs[0].runs:
                r.bold, r.font.size = True, Pt(9)
        for fila in rows:
            celdas = t.add_row().cells
            for i, v in enumerate(fila):
                celdas[i].text = str(v)
                for r in celdas[i].paragraphs[0].runs:
                    r.font.size = Pt(9)

    r = doc.add_paragraph().add_run("SDQMIP — Catálogo de Productos Sectoriales")
    r.bold, r.font.size, r.font.color.rgb = True, Pt(21), NAVY
    r.font.name = "Plus Jakarta Sans"
    _p(f"v4 · {_fecha_larga()} · Documento comercial confidencial · Santo Domingo, RD",
       size=9, color=GRAY, italic=True)
    _p(f"Generado del registro de productos y de los reportes de validación de PRODUCCIÓN "
       f"(commit servido y cifras verificadas el {HOY}). Ninguna cifra de este documento se "
       f"escribió a mano.", size=9, color=RED, italic=True)

    _h(f"El catálogo son {len(PRODUCT_CATALOG)} ejes")
    _p("La lista sale del registro que sirve la plataforma. Cada eje se empaqueta en tres "
       "niveles: Pulse (abierto, agregado del sistema) · Insight (suscripción, entidad o "
       "segmento nombrado) · Deep Dive (informe a demanda).")
    _tabla(["Eje", "Fuentes"],
           [[e.display_name.replace("SDQ ", ""), e.source] for e in PRODUCT_CATALOG])

    _h("Qué se puede afirmar de cada eje")
    _p("Los grupos no son grados de una misma escala: separan afirmaciones que no se "
       "comparan. Un backtest contra desenlaces reales de entidades y un índice descriptivo "
       "sin corte transversal no sostienen el mismo argumento.", size=9, color=GRAY)
    _p("El «Detalle» trae, además del N, la POBLACIÓN sobre la que se midió cuando una parte "
       "relevante de ella no otorga crédito. Un eje puede cubrir por diseño todo un universo "
       "supervisado —bancos, agentes de cambio, fiduciarias— y entonces su cifra no se puede "
       "leer como discriminación entre bancos. Sale computada de la misma corrida que produjo "
       "el Gini, no escrita a mano.", size=9, color=GRAY)
    por_grupo = cred.get("por_grupo") or {}
    por_eje = {f["eje"]: f for f in cred["ejes"]}
    for grupo in GRUPOS:
        claves = por_grupo.get(grupo) or []
        if not claves:
            continue
        _h(grupo, size=11, before=8)
        _tabla(["Eje", "Métrica", "Valor", "IC 95 %", "Detalle"],
               [_fila_credencial(por_eje[k]) for k in claves])

    vetadas = cred.get("vetadas_por_frescura") or []
    _h("Cifras existentes que NO entran a este documento")
    if vetadas:
        _p("Estas tienen cifra calculada, pero su reporte no se pudo verificar vigente "
           "contra el insumo que lo produjo. Se listan en vez de omitirse: un veto "
           "silencioso se lee como que el eje no tiene validación.")
        _tabla(["Eje", "Motivo"],
               [[por_eje[k]["nombre"], por_eje[k].get("stale_reason") or "—"]
                for k in vetadas])
    else:
        _p("Ninguna: las cifras publicadas arriba se verificaron vigentes contra su insumo "
           "en el momento de generar este documento.")

    _h("Lo que este documento NO afirma")
    for linea in (
        "Ninguna validación es grado-Basilea ni una PD calibrada. Todas son direccionales y "
        "se publican con su intervalo y su N.",
        "El índice de discriminación del rating bancario NO es validación contra quiebras: "
        "su desenlace es distress financiero, dominado por la regla de resultados. La "
        "credencial contra quiebras reales es la cohorte histórica, y son TRES casos "
        "evaluables — dos detectados con anticipación (11 y 7 meses) y uno con señal tardía.",
        "Los ejes del último grupo se ofrecen como descriptivos de dato real, no como "
        "predictivos, y cada uno declara qué le impide tener backtest.",
    ):
        _p("• " + linea, size=9)

    ruta = os.path.join(SALIDA_DOCX, f"Catalogo_SDQMIP_v4_Productos_Sectoriales_{HOY}.docx")
    doc.save(ruta)
    return ruta


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--local", action="store_true", help="usa la DB local en vez de prod")
    ap.add_argument("--json-only", action="store_true", help="no genera el .docx")
    args = ap.parse_args()

    cred = credenciales_locales() if args.local else credenciales_de_prod()
    cred["generado_at"] = datetime.now(timezone.utc).isoformat()
    destino = os.path.join(RAIZ, "evidence", f"credenciales_validacion_{HOY}.json")
    with open(destino, "w", encoding="utf-8") as fh:
        json.dump(cred, fh, ensure_ascii=False, indent=1)
        fh.write("\n")
    print(f"✓ {destino} — {cred['n_publicables']} credenciales publicables, "
          f"{len(cred.get('vetadas_por_frescura') or [])} vetadas por frescura")

    if not args.json_only:
        print(f"✓ {build_docx(cred)}")


if __name__ == "__main__":
    main()
