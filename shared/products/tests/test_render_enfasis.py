"""Markdown en línea en las DOS rutas de render: PDF y Word.

La cursiva de un asterisco no se convertía y salía literal en un documento que se vende: el
cliente leía `*combined ratio*`, `*loss ratio*` y el descargo de una sección entera envuelto
en asteriscos. La negrita sí funcionaba, así que el defecto sobrevivió a toda revisión que
mirara títulos.
"""
import re

from shared.products.render import _inline


class TestPDF:
    def test_la_cursiva_se_convierte(self):
        assert _inline("margen 1 − *combined ratio* = 0,28") == (
            "margen 1 − <i>combined ratio</i> = 0,28")

    def test_la_negrita_sigue_funcionando(self):
        assert _inline("**Liquidez:** el rezago") == "<b>Liquidez:</b> el rezago"

    def test_negrita_y_cursiva_conviven(self):
        """La negrita se resuelve primero: si no, `**x**` lo consume la regla de un asterisco."""
        assert _inline("**Nota:** el *loss ratio* sube") == (
            "<b>Nota:</b> el <i>loss ratio</i> sube")

    def test_no_queda_ningun_asterisco_suelto(self):
        txt = ("*El ISF es una medida de solidez por bandas elaborada por SDQ sobre estados "
               "auditados de la SIS. No es un rating de crédito.*")
        assert "*" not in _inline(txt)

    def test_un_asterisco_solitario_no_inventa_cursiva(self):
        """Un asterisco sin pareja no debe abrir una etiqueta que nunca cierra."""
        out = _inline("crecimiento de 5% * ajuste")
        assert "<i>" not in out

    def test_no_cruza_lineas(self):
        out = _inline("primer *renglón\nsegundo* renglón")
        assert "<i>" not in out

    def test_el_escapado_de_html_se_conserva(self):
        """El énfasis no debe abrir una vía para inyectar marcado."""
        assert _inline("a <b>x</b> b") == "a &lt;b&gt;x&lt;/b&gt; b"


class TestWord:
    def _runs(self, texto):
        from docx import Document

        from shared.products.render_docx import _add_runs
        p = Document().add_paragraph()
        _add_runs(p, texto)
        return [(r.text, bool(r.bold), bool(r.italic)) for r in p.runs]

    def test_cursiva_y_negrita_en_word(self):
        runs = self._runs("**Nota:** el *loss ratio* sube")
        assert ("Nota:", True, False) in runs
        assert ("loss ratio", False, True) in runs

    def test_word_no_deja_asteriscos(self):
        assert not any("*" in t for t, _, _ in self._runs("el *combined ratio* de 72%"))

    def test_las_dos_rutas_coinciden_en_que_marcan(self):
        """PDF y Word no pueden divergir en qué consideran énfasis."""
        txt = "**A** y *B* y C"
        pdf = _inline(txt)
        word = {t: (b, i) for t, b, i in self._runs(txt)}
        assert ("<b>A</b>" in pdf) and word.get("A") == (True, False)
        assert ("<i>B</i>" in pdf) and word.get("B") == (False, True)


def test_la_frescura_concuerda_el_plural():
    from shared.products.report_sections import _methodology_md

    class _Sig:
        sources, cadence, coverage, detail = ("SIS",), "quarterly", None, None
        freshness_days = 1

    # SIN corte: es el único camino que imprime la ANTIGÜEDAD en días, y por tanto el único
    # donde el plural existe. Con corte la frescura se ancla al corte y deja de contar días
    # desde hoy —porque contra hoy la cifra envejecía sola dentro de un documento fechado—;
    # ese lado lo cubre `test_frescura_anclada_al_corte`.
    md = _methodology_md(_Sig(), None, as_of=None)
    assert "un día." in md and "1 días" not in md
    _Sig.freshness_days = 5
    assert "5 días." in _methodology_md(_Sig(), None, as_of=None)
    _Sig.freshness_days = 0
    assert "menos de un día." in _methodology_md(_Sig(), None, as_of=None)


# ── Cursiva ANIDADA dentro de negrita ─────────────────────────────────────────
#
# La primera versión del arreglo compartido rompía este caso: `**negrita con *cursiva*
# adentro**` salía como `*<i>negrita con </i>cursiva<i> adentro</i>*` —negrita perdida y
# asteriscos a la vista—. Banca YA lo tenía resuelto en su propio generador de PDF (su
# docstring documentaba el mismo defecto), así que el guard estaba en un motor y faltaba en el
# compartido. Ahora hay UNA implementación y banca la consume.

_ANIDADO = "**negrita con *cursiva* adentro**"


def test_pdf_respeta_la_cursiva_dentro_de_negrita():
    assert _inline(_ANIDADO) == "<b>negrita con <i>cursiva</i> adentro</b>"


def test_no_confunde_una_multiplicacion_con_cursiva():
    """`5*8` no abre énfasis: la frontera exige que el contenido no arranque con espacio."""
    assert _inline("5*8 = 40 y 3*4 = 12") == "5*8 = 40 y 3*4 = 12"


def test_word_respeta_la_cursiva_dentro_de_negrita():
    from docx import Document

    from shared.products.render_docx import _add_runs
    p = Document().add_paragraph()
    _add_runs(p, _ANIDADO)
    runs = [(r.text, bool(r.bold), bool(r.italic)) for r in p.runs]
    assert ("cursiva", True, True) in runs, runs
    assert all("*" not in tx for tx, _, _ in runs)


def test_banca_no_mantiene_una_segunda_implementacion():
    """Dos copias del mismo criterio divergen; banca delega en la compartida."""
    from modules.banking_score.reports.pdf_generator import _md_inline
    assert _md_inline(_ANIDADO) == _inline(_ANIDADO)
