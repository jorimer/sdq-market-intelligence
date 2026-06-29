"""Tests del renderer Word (.docx) de marca y del dispatch por formato."""
import os

from docx import Document

from shared.products.render import render_product_pdf
from shared.products.render_docx import render_product_docx

_NARR = {
    "assessment": "## Lectura\n> ITT 93.0 — demanda récord.\nLa **construcción** resta.\n- Comercio +0.26pp",
    "std_methodology": "**Fuentes:** BCRD. **Cobertura:** 100%.",
}


def test_docx_renders_with_brand_anatomy(tmp_path):
    path = render_product_docx(
        sector_key="demo", display_name="Economía Dominicana", title="Deep Dive Estructura",
        period="2025", headline="VAB +2.03% — lo mueven los servicios.",
        narratives=_NARR, section_titles={"assessment": "Evaluación", "std_methodology": "Metodología y fuentes"},
        tables=[("Dimensiones", [["Dim", "Score"], ["financiero", "100"]])],
        watermark="On-demand · SDQMIP", output_dir=str(tmp_path))
    assert os.path.exists(path) and path.endswith(".docx")
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    # texto de las celdas de tabla (la banda de portada y la tabla de datos son tablas)
    cells = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    full = text + "\n" + cells
    assert "Deep Dive Estructura" in full            # título (en la banda de portada)
    assert "Economía Dominicana" in text              # sujeto
    assert "VAB +2.03%" in text                        # pull-quote de portada
    assert "1.  Evaluación" in text                    # sección numerada
    assert "Metodología y fuentes" in text             # sección estándar
    assert len(doc.inline_shapes) >= 1                 # logo embebido
    assert "financiero" in cells                        # tabla de datos


def test_render_product_pdf_dispatches_to_docx(tmp_path):
    path = render_product_pdf(
        sector_key="demo", display_name="X", title="Insight X", period="2025",
        narratives={"a": "texto"}, output_dir=str(tmp_path), fmt="docx")
    assert path.endswith(".docx") and os.path.exists(path)


def test_pdf_still_default(tmp_path):
    path = render_product_pdf(
        sector_key="demo", display_name="X", title="Insight X", period="2025",
        narratives={"a": "texto"}, output_dir=str(tmp_path))
    assert path.endswith(".pdf") and os.path.exists(path)
