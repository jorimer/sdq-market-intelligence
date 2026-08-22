"""Renderer Word (.docx) de marca — misma anatomía que el PDF (docs/REPORT_STANDARD.md).

Mismo contrato de entrada que ``render_product_pdf`` (sector_key, display_name, title,
period, narratives, section_titles, tables, subtitle, headline, watermark, sample) → un
.docx editable por el cliente con la marca SDQ·MIP: portada con logo + banda navy + título,
secciones numeradas con pull-quotes, tablas, encabezado corrido + número de página, y la
estampa de muestra / watermark al pie. No duplica lógica de dominio; pinta el contenido.
"""
from __future__ import annotations

import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from shared import brand
from shared.config.settings import settings

def _rgb(token: str) -> RGBColor:
    """Token de marca (``#RRGGBB``) → color de python-docx."""
    return RGBColor.from_string(token.lstrip("#").upper())


def _hex(token: str) -> str:
    """Token de marca → hex sin almohadilla, que es como lo quiere el sombreado OOXML."""
    return token.lstrip("#").upper()


# Paleta: se LEE de `shared.brand`. Ver el comentario equivalente en `render.py`.
_NAVY = _rgb(brand.INK)            # tinta: títulos y cabeceras
_BLUE = _rgb(brand.ACCENT_INK)     # texto en acento — el que contrasta sobre blanco
_ACCENT = _hex(brand.ACCENT)       # rellenos de acento: barra del pull-quote
_GRAY = _rgb(brand.MUTED)
_ALERT = _rgb(brand.ALERT)         # estampa de MUESTRA
_WHITE = _rgb(brand.WHITE)
_NAVY_HEX = _hex(brand.INK)

_LOGO = os.path.join(os.path.dirname(__file__), "assets", "sdq_mip_logo.png")
_GLYPH_RE = re.compile("[▀-▟■-◿✀-➿☀-⛿\U0001f000-\U0001ffff️]")

DISCLAIMER_ES = (
    "Las opiniones expresadas en este informe son de SDQ Consulting y no constituyen "
    "una recomendación de inversión. SDQ no asume responsabilidad por pérdidas "
    "derivadas del uso de esta información."
)


def _clean(text: str) -> str:
    return _GLYPH_RE.sub("", text or "").strip()


def _shade(cell, hex_color: str) -> None:
    """Sombrea una celda de tabla con un color de fondo (vía XML)."""
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(sh)


def _left_accent(paragraph, hex_color: str) -> None:
    """Barra de acento a la izquierda del párrafo (pull-quote)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), hex_color)
    pbdr.append(left)
    pPr.append(pbdr)


def _add_runs(paragraph, text: str, *, color: Optional[RGBColor] = None, size: Optional[int] = None,
              bold_all: bool = False) -> None:
    """Añade *text* a *paragraph* respetando **negritas** y *cursivas* markdown.

    La cursiva faltaba y salía literal, igual que en el PDF: el lector veía los asteriscos.
    Se parte primero por `**` —los impares son negrita— y dentro de cada trozo restante por
    `*`, de modo que `**x**` no lo consuma la regla de un asterisco.
    """
    from shared.products.render import _ITALIC_RE

    def _emitir(txt: str, *, negrita: bool) -> None:
        """Parte por CURSIVA y emite. Un mismo criterio de frontera que el PDF."""
        pos = 0
        for m in _ITALIC_RE.finditer(txt):
            antes = txt[pos:m.start()] + m.group(1)
            if antes:
                r = paragraph.add_run(antes)
                r.bold = negrita or bold_all
                _estilo(r, color, size)
            r = paragraph.add_run(m.group(2))
            r.bold = negrita or bold_all
            r.italic = True                  # cursiva, anidada dentro de negrita si toca
            _estilo(r, color, size)
            pos = m.end()
        resto = txt[pos:]
        if resto:
            r = paragraph.add_run(resto)
            r.bold = negrita or bold_all
            _estilo(r, color, size)

    # Los impares del split vienen de **...** — y pueden contener *cursiva* adentro.
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", _clean(text))):
        if chunk:
            _emitir(chunk, negrita=i % 2 == 1)


def _estilo(run, color: Optional[RGBColor], size: Optional[int]) -> None:
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def _page_number_field(paragraph) -> None:
    """Inserta el campo PAGE (número de página) en un párrafo de pie."""
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addprevious(fld)


def _furniture(doc, header_line: str, watermark: Optional[str], sample: bool) -> None:
    """Encabezado corrido + pie con watermark/estampa y número de página."""
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hr = hp.add_run(header_line)
    hr.font.size = Pt(8)
    hr.font.bold = True
    hr.font.color.rgb = _NAVY
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wm = "MUESTRA — DATA ILUSTRATIVA" if sample else (watermark or "")
    if wm:
        fr = fp.add_run(wm + "      ")
        fr.font.size = Pt(8)
        fr.font.color.rgb = _ALERT if sample else _GRAY
    _page_number_field(fp)


def _md_split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _md_is_sep(line: str) -> bool:
    """Fila separadora de tabla markdown: `|---|:--:|---|`."""
    t = line.strip()
    return "-" in t and "|" in t and all(
        re.fullmatch(r":?-{1,}:?", c) for c in _md_split_row(t))


def _md_table(doc, header: List[str], rows: List[List[str]]) -> None:
    """Tabla markdown → tabla branded (mismo look que las tablas de datos del Word)."""
    ncol = len(header) or 1
    data = [header] + [(r + [""] * ncol)[:ncol] for r in rows]
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(data):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = str(row[ci]) if ci < len(row) else ""
            _add_runs(cells[ci].paragraphs[0], val,
                      color=(_WHITE if ri == 0 else None), bold_all=(ri == 0))
            if ri == 0:
                _shade(cells[ci], _NAVY_HEX)


def _md_body(doc, text: str) -> None:
    """Cuerpo markdown-lite: tabla · --- regla · ## subhead · > pull-quote · -/* viñeta · **negrita**."""
    lines = (text or "").replace("\r", "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        # Tabla markdown: fila de encabezado + fila separadora + cuerpo.
        if "|" in line and i + 1 < len(lines) and _md_is_sep(lines[i + 1]):
            header = _md_split_row(line)
            body: List[List[str]] = []
            j = i + 2
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                body.append(_md_split_row(lines[j]))
                j += 1
            _md_table(doc, header, body)
            i = j
            continue
        # Regla horizontal (`---`): separador de párrafo, no texto literal.
        if re.fullmatch(r"-{3,}", line):
            doc.add_paragraph()
            i += 1
            continue
        q = re.match(r"^>\s+(.*)$", line)
        h = re.match(r"^(#{1,3})\s+(.*)$", line)
        if q:
            p = doc.add_paragraph()
            _left_accent(p, _ACCENT)
            _add_runs(p, q.group(1), color=_NAVY, size=13)
        elif h:
            p = doc.add_paragraph()
            _add_runs(p, h.group(2), color=_BLUE, size=12, bold_all=True)
        elif re.match(r"^(?:[-*]|\d+[.)])\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^(?:[-*]|\d+[.)])\s+", "", line))
        else:
            _add_runs(doc.add_paragraph(), line)
        i += 1


def render_product_docx(
    *,
    sector_key: str,
    display_name: str,
    title: str,
    period: str,
    narratives: Dict[str, str],
    section_titles: Optional[Dict[str, str]] = None,
    tables: Optional[List[Tuple[str, Sequence[Sequence[str]]]]] = None,
    charts: Optional[List[dict]] = None,
    subtitle: Optional[str] = None,
    headline: Optional[str] = None,
    watermark: Optional[str] = None,
    sample: bool = False,
    output_dir: Optional[str] = None,
    tables_last: bool = False,
) -> str:
    """Renderiza el .docx de marca de un producto y devuelve el path."""
    from shared.products.report_sections import STANDARD_SECTION_TITLES
    section_titles = {**STANDARD_SECTION_TITLES, **(section_titles or {})}

    out_dir = output_dir or settings.REPORTS_DIR
    os.makedirs(out_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = re.sub(r"[^a-z0-9]+", "_", display_name.lower()).strip("_")
    path = os.path.join(out_dir, f"{sector_key}_{safe}_{ts}.docx")

    doc = Document()
    _furniture(doc, f"SDQ·MIP — {title} · {display_name}", watermark, sample)

    # ── Portada ──
    if os.path.exists(_LOGO):
        doc.add_picture(_LOGO, width=Inches(0.5))
    kicker = doc.add_paragraph()
    _add_runs(kicker, "SDQ·MIP · MARKET INTELLIGENCE REPORT", color=_BLUE, size=10, bold_all=True)
    band = doc.add_table(rows=1, cols=1)
    cell = band.rows[0].cells[0]
    _shade(cell, _NAVY_HEX)
    cp = cell.paragraphs[0]
    _add_runs(cp, title, color=_WHITE, size=22, bold_all=True)
    sub = doc.add_paragraph()
    _add_runs(sub, display_name, color=_NAVY, size=16, bold_all=True)
    if subtitle:
        _add_runs(doc.add_paragraph(), subtitle, color=_BLUE, size=12)
    if headline:
        hq = doc.add_paragraph()
        _left_accent(hq, _ACCENT)
        _add_runs(hq, headline, color=_NAVY, size=13)
    _add_runs(doc.add_paragraph(), f"**Período:** {period}", color=_NAVY)
    _add_runs(doc.add_paragraph(), f"**Fecha:** {datetime.now().strftime('%d/%m/%Y')}", color=_NAVY)
    doc.add_page_break()

    # ── Gráficos de marca (barras) ──
    from shared.products.charts import render_charts
    for _ctitle, png in render_charts(charts, out_dir, f"{sector_key}_{ts}"):
        doc.add_picture(png, width=Inches(6.0))

    # ── Tablas de datos ── (con tables_last van DESPUÉS del texto, como anexo)
    def _emit_tables() -> None:
        for heading, rows in (tables or []):
            if not rows:
                continue
            _add_runs(doc.add_paragraph(), heading, color=_NAVY, size=15, bold_all=True)
            ncol = max((len(r) for r in rows), default=1)
            t = doc.add_table(rows=0, cols=ncol)
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci in range(ncol):
                    val = str(row[ci]) if ci < len(row) else ""
                    _add_runs(cells[ci].paragraphs[0], val,
                              color=(_WHITE if ri == 0 else None), bold_all=(ri == 0))
                    if ri == 0:
                        _shade(cells[ci], _NAVY_HEX)

    if not tables_last:
        _emit_tables()

    # ── Secciones narrativas numeradas ──
    for n, (key, text) in enumerate(narratives.items(), start=1):
        ttl = section_titles.get(key, key.replace("_", " ").title())
        _add_runs(doc.add_paragraph(), f"{n}.  {ttl}", color=_NAVY, size=15, bold_all=True)
        _md_body(doc, text)

    if tables_last and tables:
        _add_runs(doc.add_paragraph(), "Anexo de datos", color=_BLUE, size=13, bold_all=True)
        _emit_tables()

    _add_runs(doc.add_paragraph(), "Disclaimer", color=_BLUE, size=12, bold_all=True)
    _add_runs(doc.add_paragraph(), DISCLAIMER_ES, color=_GRAY, size=8)

    doc.save(path)
    return path
