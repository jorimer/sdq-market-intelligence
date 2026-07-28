"""Informe forense en Word (.docx) branded — paridad de FORMA con el PDF forense.

Misma anatomía y misma lógica (`forensic_common`) que `render_forensic_pdf`: portada de marca,
stat cards, los tres gráficos con marcas de onset/colapso y mediana de pares, timeline de la
lectura del modelo, dictamen de legibilidad y fuentes — pero **editable**. Reusa los gráficos
matplotlib del PDF forense y los helpers de marca del renderer docx genérico. Devuelve bytes.
"""
from __future__ import annotations

import os
import tempfile
from io import BytesIO
from typing import Dict

from docx import Document
from docx.shared import Inches, RGBColor

from modules.banking_score.reports.forensic_common import (
    humanize_month,
    legibility,
    marker_indices,
    model_timeline,
    stat_cards,
)
from modules.banking_score.reports.forensic_pdf import (
    _chart_cobertura,
    _chart_deposito,
    _chart_morosidad,
)
from shared.products.render_docx import (
    DISCLAIMER_ES,
    _BLUE,
    _GRAY,
    _LOGO,
    _NAVY,
    _NAVY_HEX,
    _SIGNAL,
    _WHITE,
    _add_runs,
    _furniture,
    _left_accent,
    _md_body,
    _md_table,
    _shade,
)

_ALERT = RGBColor(0xC8, 0x39, 0x2E)
_WARN = RGBColor(0xB7, 0x79, 0x1F)
_OK = RGBColor(0x15, 0x87, 0x5A)
_MUTED = RGBColor(0x76, 0x82, 0x9C)
_TONE = {"alert": _ALERT, "warn": _WARN, "ok": _OK, "ink": _NAVY}


def render_forensic_docx(pkg: Dict, narrative_md: str, *, degraded: bool = False) -> bytes:
    """Documento forense en Word (bytes) — portada branded + cuerpo con paridad al PDF."""
    from modules.banking_score.historical_service import forensic_narrative_context

    meta, btk, series = pkg["meta"], pkg["backtest"], pkg["series"]
    ctx = forensic_narrative_context(pkg)
    marks = marker_indices(pkg)
    lead = btk.get("lead_months")
    onset = btk.get("onset_cluster")

    doc = Document()
    _furniture(doc, f"SDQ·MIP — Informe Forense · {meta['nombre']}", None, False)

    # ── Portada de marca ──
    if os.path.exists(_LOGO):
        doc.add_picture(_LOGO, width=Inches(0.5))
    _add_runs(doc.add_paragraph(), "SDQ·MIP — INFORME FORENSE · RETROSPECTIVO",
              color=_BLUE, size=10, bold_all=True)
    band = doc.add_table(rows=1, cols=1)
    cell = band.rows[0].cells[0]
    _shade(cell, _NAVY_HEX)
    _add_runs(cell.paragraphs[0], meta["nombre"], color=_WHITE, size=22, bold_all=True)
    _add_runs(doc.add_paragraph(),
              "Anatomía de una quiebra · reconstrucción del deterioro sobre dato mensual real",
              color=_BLUE, size=12)
    _add_runs(doc.add_paragraph(),
              f"**Período revisado:** {humanize_month(meta['primer'])} → {humanize_month(meta['ultimo'])}",
              color=_NAVY, size=10)
    verdict = (f"Deterioro detectable desde {humanize_month(onset)} — {lead} meses antes del "
               f"colapso." if onset and lead is not None
               else "Los ratios reportados nunca formaron un cluster de alerta antes de la salida.")
    vq = doc.add_paragraph()
    _left_accent(vq, _SIGNAL)
    _add_runs(vq, verdict, color=_NAVY, size=13)

    # ── Stat cards (4 cifras) ──
    _add_runs(doc.add_paragraph(), "Resumen ejecutivo", color=_NAVY, size=15, bold_all=True)
    cards = stat_cards(pkg, ctx)
    ct = doc.add_table(rows=1, cols=4)
    ct.style = "Light Grid Accent 1"
    for cell, c in zip(ct.rows[0].cells, cards):
        _add_runs(cell.paragraphs[0], c["value"], color=_TONE.get(c["tone"], _NAVY), size=16,
                  bold_all=True)
        _add_runs(cell.add_paragraph(), c["label"], color=_MUTED, size=8)

    # ── Gráficos + timeline + legibilidad + narrativa + fuentes ──
    tmp = tempfile.mkdtemp(prefix="forensic_docx_")
    try:
        cm, cc, cd = (os.path.join(tmp, n) for n in ("mora.png", "cob.png", "dep.png"))
        _chart_morosidad(series, marks, cm)
        _chart_cobertura(series, marks, cc)
        _chart_deposito(series, marks, cd)
        for titulo, png in (("Riesgo de crédito — la morosidad se despega de sus pares", cm),
                            ("Colchón de provisiones — la cobertura se agota", cc),
                            ("Fuga de depósitos (variación intermensual)", cd)):
            _add_runs(doc.add_paragraph(), titulo, color=_NAVY, size=14, bold_all=True)
            doc.add_picture(png, width=Inches(6.2))

        _add_runs(doc.add_paragraph(), "Qué habría marcado Alerta Temprana, y cuándo",
                  color=_NAVY, size=14, bold_all=True)
        for r in model_timeline(pkg, ctx):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, f"[{r['flag']}] ", color=_TONE.get(r["tone"], _WARN), bold_all=True)
            _add_runs(p, f"{r['fecha']} — ", color=_NAVY, bold_all=True)
            _add_runs(p, r["text"])

        leg = legibility(pkg, ctx)
        lbox = doc.add_table(rows=1, cols=1)
        lcell = lbox.rows[0].cells[0]
        _shade(lcell, "E5F3EC" if leg["legible"] else "FBF1DD")
        _add_runs(lcell.paragraphs[0], leg["title"], color=(_OK if leg["legible"] else _WARN),
                  size=12, bold_all=True)
        _add_runs(lcell.add_paragraph(), leg["text"], color=_NAVY, size=10)

        _add_runs(doc.add_paragraph(), "Lectura forense", color=_NAVY, size=15, bold_all=True)
        if narrative_md and not degraded:
            _md_body(doc, narrative_md)
        else:
            _add_runs(doc.add_paragraph(),
                      "La lectura narrativa no está disponible en este momento; los datos y el "
                      "backtest de arriba son completos.")

        _add_runs(doc.add_paragraph(), "Metodología y fuentes", color=_NAVY, size=15, bold_all=True)
        _add_runs(doc.add_paragraph(),
                  "Todas las cifras salen del estado de situación y de resultados mensual por "
                  "entidad publicado por la Superintendencia de Bancos (Cronología SB). El inicio "
                  "del deterioro es el primer mes con un cluster de ≥2 alertas altas simultáneas; "
                  "la mediana de pares es la morosidad mensual de las entidades del mismo tipo.")
        _md_table(doc, ["Fuente", "Serie", "Cobertura"], [
            ["SB — Cronología SB", "Estado de situación mensual, por entidad",
             f"{humanize_month(meta['primer'])} → {humanize_month(meta['ultimo'])}"],
            ["SB — Cronología SB", "Estado de resultados mensual, por entidad", "ene 1996 → abr 2026"],
        ])
        _add_runs(doc.add_paragraph(),
                  "Límite: el capital regulatorio Basilea no existe en el balance contable "
                  "pre-2004 — el apalancamiento patrimonio/activos es un proxy etiquetado. Informe "
                  "retrospectivo/forense, no una calificación emitida en su momento.",
                  color=_GRAY, size=8)

        _add_runs(doc.add_paragraph(), "Disclaimer", color=_BLUE, size=12, bold_all=True)
        _add_runs(doc.add_paragraph(), DISCLAIMER_ES, color=_GRAY, size=8)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    finally:
        for n in ("mora.png", "cob.png", "dep.png"):
            png_path = os.path.join(tmp, n)
            if os.path.exists(png_path):
                os.remove(png_path)
        os.rmdir(tmp)
